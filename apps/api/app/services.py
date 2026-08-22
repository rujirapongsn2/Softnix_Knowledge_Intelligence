import hashlib
import logging
import mimetypes
import re
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta
from pathlib import Path
from threading import BoundedSemaphore

from bs4 import BeautifulSoup
from docx import Document as WordDocument
import httpx
from markitdown import MarkItDown
from sqlalchemy import Text, and_, cast, func, or_
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session, aliased

from .config import get_settings
from .db import SessionLocal
from .graph_store import Neo4jGraphStore
from .legal_registry import AUTHORITY_LEVELS, classify_kind, normalize_family_key, parse_provision_refs, parse_thai_date, provision_number_matches, resolve_instrument_statuses
from .legal_resolver import resolve_legal_context
from .legal_corpus import parse_legal_corpus_metadata
from .document_templates import metadata_search_text, normalize_field_definitions
from .models import Document, DocumentChunk, DocumentMetadataValue, Entity, EntitySource, GraphProjectionEvent, KnowledgeBase, LegalFamily, LegalInstrument, LegalInstrumentRelation, ProcessingJob, QueryResult, Relationship, RelationshipSource
from .openrouter import OpenRouterClient
from .observability import metrics
from .planner import LegalContext, RetrievalChannel, RetrievalPlan, PlannerDecision, apply_llm_plan, intersect_policies, policy_from_config, rule_plan
from .retrieval import LightRAGRetrievalEngine, RetrievalEvidence

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".pptx", ".xlsx", ".xls", ".txt", ".md", ".html", ".htm", ".csv", ".json"}
LEGACY_EXTRACTOR_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm", ".csv", ".json"}
ALLOWED_MIME_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    ".pptx": {"application/vnd.openxmlformats-officedocument.presentationml.presentation"},
    ".xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
    ".xls": {"application/vnd.ms-excel"},
    ".txt": {"text/plain"}, ".md": {"text/markdown", "text/plain"},
    ".html": {"text/html"}, ".htm": {"text/html"},
    ".csv": {"text/csv", "application/csv", "text/plain"},
    ".json": {"application/json", "text/json", "text/plain"},
}
DEFAULT_RETRIEVAL_CONFIG = {
    "version": 1, "retrieval_mode": "auto", "enable_vector": True, "enable_graph": True,
    "enable_fulltext": True, "enable_lightrag": True, "enable_reranker": True,
    "planner_llm_fallback": True, "default_top_k": 12, "maximum_top_k": 30,
    "maximum_graph_depth": 3, "citation_required": True,
}
TRANSIENT_PROCESSING_ERRORS = {
    "RETRIEVAL_ENGINE_UNAVAILABLE", "RETRIEVAL_ENGINE_REJECTED", "RETRIEVAL_ENGINE_BUSY",
    "RETRIEVAL_ENGINE_TIMEOUT", "RETRIEVAL_ENGINE_BUDGET_EXHAUSTED", "OPENROUTER_UNAVAILABLE", "EXTERNAL_OCR_UNAVAILABLE", "EXTERNAL_OCR_TIMEOUT", "OCR_CHAIN_FAILED",
}
MAX_PROCESSING_ATTEMPTS = 3
# Budget exhaustion outlives the normal retry window (the shared provider
# account frees in-flight quota on the scale of minutes), so those tracks
# earn extra attempts instead of failing the document permanently (F1).
MAX_BUDGET_PROCESSING_ATTEMPTS = 6
# Cap on how many documents one worker may push into the remote indexing
# pipeline at the same time.  LightRAG fans every insert out into upstream LLM
# calls; 13 parallel ingests exhausted the shared OpenRouter in-flight budget
# (402) and failed the whole batch (F1).
INDEXING_CONCURRENCY_LIMIT = 2
# A budget-exhausted track is retried on the same exponential schedule as the
# other transient errors, but with a longer floor: the account's in-flight
# quota frees up on the scale of minutes, not seconds.
BUDGET_RETRY_DELAY_FLOOR_SECONDS = 60
_indexing_semaphore = BoundedSemaphore(INDEXING_CONCURRENCY_LIMIT)
logger = logging.getLogger(__name__)


def processing_retry_delay(attempt_count: int) -> int:
    return min(60, 2 ** max(1, attempt_count))


def canonical_entity_name(value: str) -> str:
    return " ".join(value.casefold().split())


def generic_identity_key(name: str) -> str:
    return f"generic:{canonical_entity_name(name)}"


def create_entity(db: Session, knowledge_base_id: str, payload) -> Entity:
    canonical = canonical_entity_name(payload.name)
    identity_key = generic_identity_key(payload.name)
    entity = db.query(Entity).filter_by(knowledge_base_id=knowledge_base_id, identity_key=identity_key, entity_type=payload.entity_type).first()
    if not entity:
        entity = Entity(knowledge_base_id=knowledge_base_id, name=payload.name, canonical_name=canonical,
                        entity_type=payload.entity_type, description=payload.description, aliases=payload.aliases,
                        attributes=payload.attributes, confidence=payload.confidence, identity_key=identity_key,
                        origin="manual", review_status="verified", is_legal=False)
        db.add(entity)
        db.flush()
    elif entity.deleted_at:
        entity.deleted_at, entity.name, entity.description = None, payload.name, payload.description
    if payload.document_id and payload.excerpt:
        exists = db.query(EntitySource).filter_by(entity_id=entity.id, document_id=payload.document_id, excerpt=payload.excerpt).first()
        if not exists:
            db.add(EntitySource(entity_id=entity.id, document_id=payload.document_id, excerpt=payload.excerpt))
            entity.source_count += 1
    db.add(GraphProjectionEvent(event_type="entity", entity_id=entity.id))
    db.commit()
    db.refresh(entity)
    return entity


def create_relationship(db: Session, knowledge_base_id: str, payload) -> Relationship:
    source = db.get(Entity, payload.source_entity_id)
    target = db.get(Entity, payload.target_entity_id)
    if not source or not target or source.knowledge_base_id != knowledge_base_id or target.knowledge_base_id != knowledge_base_id:
        raise ValueError("ENTITY_NOT_FOUND")
    relationship = db.query(Relationship).filter_by(knowledge_base_id=knowledge_base_id, source_entity_id=source.id,
                                                     target_entity_id=target.id, relationship_type=payload.relationship_type).first()
    if not relationship:
        relationship = Relationship(knowledge_base_id=knowledge_base_id, source_entity_id=source.id, target_entity_id=target.id,
                                    relationship_type=payload.relationship_type, description=payload.description, confidence=payload.confidence,
                                    origin="manual", review_status="verified", is_legal=False)
        db.add(relationship)
        db.flush()
    elif relationship.deleted_at:
        relationship.deleted_at, relationship.description, relationship.confidence = None, payload.description, payload.confidence
    if payload.document_id and payload.excerpt:
        exists = db.query(RelationshipSource).filter_by(relationship_id=relationship.id, document_id=payload.document_id, excerpt=payload.excerpt).first()
        if not exists:
            db.add(RelationshipSource(relationship_id=relationship.id, document_id=payload.document_id, excerpt=payload.excerpt))
            relationship.source_count += 1
    db.add(GraphProjectionEvent(event_type="relationship", relationship_id=relationship.id))
    db.commit()
    db.refresh(relationship)
    return relationship


def _remove_metadata_graph_projection(db: Session, document: Document) -> None:
    """Remove only graph facts generated from the document's metadata fields."""
    relationship_ids = [row[0] for row in db.query(Relationship.id).join(RelationshipSource).filter(
        RelationshipSource.document_id == document.id, Relationship.origin == "metadata"
    ).all()]
    for relationship_id in relationship_ids:
        db.query(RelationshipSource).filter_by(relationship_id=relationship_id, document_id=document.id).delete(synchronize_session=False)
        relationship = db.get(Relationship, relationship_id)
        if relationship:
            relationship.source_count = db.query(func.count(RelationshipSource.id)).filter_by(relationship_id=relationship_id).scalar() or 0
            if relationship.source_count == 0:
                relationship.deleted_at = datetime.utcnow()
    entity_ids = [row[0] for row in db.query(Entity.id).join(EntitySource).filter(
        EntitySource.document_id == document.id, Entity.origin == "metadata"
    ).all()]
    for entity_id in entity_ids:
        db.query(EntitySource).filter_by(entity_id=entity_id, document_id=document.id).delete(synchronize_session=False)
        entity = db.get(Entity, entity_id)
        if entity:
            entity.source_count = db.query(func.count(EntitySource.id)).filter_by(entity_id=entity_id).scalar() or 0
            if entity.source_count == 0:
                entity.deleted_at = datetime.utcnow()
    db.flush()


def sync_document_metadata_values(db: Session, document: Document) -> None:
    """Maintain the indexed projection used by exact metadata filters."""
    db.query(DocumentMetadataValue).filter(DocumentMetadataValue.document_id == document.id).delete(synchronize_session=False)
    fields = normalize_field_definitions(document.metadata_template_fields)
    values = document.document_metadata or {}
    for field in fields:
        key = field.get("key")
        value = values.get(key) if key else None
        if not field.get("filterable") or value in (None, ""):
            continue
        db.add(DocumentMetadataValue(
            knowledge_base_id=document.knowledge_base_id,
            document_id=document.id,
            field_key=key,
            value_text=str(value)[:10000],
        ))
    db.flush()


def sync_document_metadata_graph(db: Session, document: Document) -> dict[str, int]:
    """Promote only explicitly mapped metadata fields into the local graph.

    The document and mapped values retain a source excerpt, so graph evidence
    remains traceable to the user-supplied metadata rather than becoming an
    unproven global fact.
    """
    _remove_metadata_graph_projection(db, document)
    fields = normalize_field_definitions(document.metadata_template_fields)
    mapped = [field for field in fields if field.get("graph_relationship") and field.get("graph_entity_type")
              and document.document_metadata.get(field.get("key")) not in (None, "")]
    if not mapped:
        return {"entities": 0, "relationships": 0}
    anchor_key = f"metadata:document:{document.id}"
    anchor = db.query(Entity).filter_by(knowledge_base_id=document.knowledge_base_id, identity_key=anchor_key, entity_type="Document").first()
    created_anchor = False
    if not anchor:
        anchor = Entity(knowledge_base_id=document.knowledge_base_id,
                        name=document.title or document.original_filename,
                        canonical_name=canonical_entity_name(document.title or document.original_filename),
                        identity_key=anchor_key, entity_type="Document", confidence=1.0,
                        origin="metadata", review_status="verified", is_legal=False,
                        attributes={"document_id": document.id})
        db.add(anchor); db.flush(); created_anchor = True
        db.add(GraphProjectionEvent(event_type="entity", entity_id=anchor.id))
    else:
        anchor.deleted_at = None
    anchor_excerpt = f"Document: {document.title or document.original_filename}"[:5000]
    if not db.query(EntitySource).filter_by(entity_id=anchor.id, document_id=document.id, excerpt=anchor_excerpt).first():
        db.add(EntitySource(entity_id=anchor.id, document_id=document.id, excerpt=anchor_excerpt)); anchor.source_count += 1
    entity_count, relationship_count = int(created_anchor), 0
    for field in mapped:
        raw_value = document.document_metadata.get(field["key"])
        value = str(raw_value).strip()[:500]
        entity_type = str(field["graph_entity_type"]).strip()[:100]
        relationship_type = str(field["graph_relationship"]).strip().upper()
        identity_key = f"metadata:{entity_type.casefold()}:{canonical_entity_name(value)}"
        target = db.query(Entity).filter_by(knowledge_base_id=document.knowledge_base_id,
                                             identity_key=identity_key, entity_type=entity_type).first()
        if not target:
            target = Entity(knowledge_base_id=document.knowledge_base_id, name=value,
                            canonical_name=canonical_entity_name(value), identity_key=identity_key,
                            entity_type=entity_type, confidence=1.0, origin="metadata",
                            review_status="verified", is_legal=False,
                            attributes={"metadata_field": field["key"]})
            db.add(target); db.flush()
            db.add(GraphProjectionEvent(event_type="entity", entity_id=target.id))
            entity_count += 1
        else:
            target.deleted_at = None
        excerpt = f"{field.get('label') or field['key']}: {value}"[:5000]
        if not db.query(EntitySource).filter_by(entity_id=target.id, document_id=document.id, excerpt=excerpt).first():
            db.add(EntitySource(entity_id=target.id, document_id=document.id, excerpt=excerpt)); target.source_count += 1
        edge = db.query(Relationship).filter_by(knowledge_base_id=document.knowledge_base_id,
                                                source_entity_id=anchor.id, target_entity_id=target.id,
                                                relationship_type=relationship_type).first()
        if not edge:
            edge = Relationship(knowledge_base_id=document.knowledge_base_id, source_entity_id=anchor.id,
                                target_entity_id=target.id, relationship_type=relationship_type,
                                description=excerpt, confidence=1.0, origin="metadata",
                                review_status="verified", is_legal=False)
            db.add(edge); db.flush(); relationship_count += 1
            db.add(GraphProjectionEvent(event_type="relationship", relationship_id=edge.id))
        else:
            edge.deleted_at, edge.description, edge.origin, edge.review_status = None, excerpt, "metadata", "verified"
        if not db.query(RelationshipSource).filter_by(relationship_id=edge.id, document_id=document.id, excerpt=excerpt).first():
            db.add(RelationshipSource(relationship_id=edge.id, document_id=document.id, excerpt=excerpt)); edge.source_count += 1
    db.flush()
    return {"entities": entity_count, "relationships": relationship_count}


def sync_lightrag_document_graph(db: Session, document: Document, max_labels: int = 200) -> dict[str, int]:
    """Materialize document-scoped LightRAG graph evidence in the editable local graph.

    LightRAG keeps one shared graph service, so file-source provenance is the
    boundary that prevents nodes from a different Knowledge Base leaking in.
    """
    # LightRAG's opaque relationship labels are useful retrieval signals but
    # are not legal facts.  Legal documents get their graph only from the
    # evidence-backed legal projection below.
    if document.document_type in LEGAL_DOCUMENT_TYPES:
        return {"entities": 0, "relationships": 0}
    engine = LightRAGRetrievalEngine()
    if not engine.enabled:
        return {"entities": 0, "relationships": 0}
    marker = f"softnix-kb={document.knowledge_base_id}__doc={document.id}__"
    graphs: list[dict] = []
    try:
        for label in engine.graph_labels()[:max_labels]:
            graph = engine.graph(label, max_depth=1, max_nodes=100)
            if any(marker in str(node.get("properties", {}).get("file_path", "")) for node in graph.get("nodes", [])):
                graphs.append(graph)
    except RuntimeError:
        # Graph sync must never mark a successfully indexed document as failed.
        return {"entities": 0, "relationships": 0}

    raw_nodes: dict[str, dict] = {}
    raw_edges: dict[str, dict] = {}
    for graph in graphs:
        for node in graph.get("nodes", []):
            props = node.get("properties", {})
            if marker in str(props.get("file_path", "")):
                raw_nodes.setdefault(str(node.get("id") or props.get("entity_id") or ""), node)
        for edge in graph.get("edges", []):
            props = edge.get("properties", {})
            if marker in str(props.get("file_path", "")):
                raw_edges.setdefault(str(edge.get("id") or f"{edge.get('source')}->{edge.get('target')}"), edge)

    local_by_external: dict[str, Entity] = {}
    entity_count = relationship_count = 0
    for external_id, node in raw_nodes.items():
        if not external_id:
            continue
        props = node.get("properties", {})
        name = str(props.get("entity_id") or node.get("labels", [external_id])[0] or external_id).strip()
        if not name:
            continue
        entity_type = str(props.get("entity_type") or "Concept").strip()[:100] or "Concept"
        canonical = canonical_entity_name(name)
        identity_key = generic_identity_key(name)
        entity = db.query(Entity).filter_by(knowledge_base_id=document.knowledge_base_id, identity_key=identity_key, entity_type=entity_type).first()
        if not entity:
            entity = Entity(knowledge_base_id=document.knowledge_base_id, name=name, canonical_name=canonical, entity_type=entity_type,
                            description=str(props.get("description") or "")[:5000] or None, confidence=1.0,
                            identity_key=identity_key, origin="ai_suggestion", review_status="suggested", is_legal=False)
            db.add(entity); db.flush(); entity_count += 1
            db.add(GraphProjectionEvent(event_type="entity", entity_id=entity.id))
        elif entity.deleted_at:
            entity.deleted_at = None
        local_by_external[external_id] = entity
        excerpt = str(props.get("description") or name)[:5000]
        if not db.query(EntitySource).filter_by(entity_id=entity.id, document_id=document.id, excerpt=excerpt).first():
            db.add(EntitySource(entity_id=entity.id, document_id=document.id, excerpt=excerpt)); entity.source_count += 1

    for edge in raw_edges.values():
        source, target = local_by_external.get(str(edge.get("source"))), local_by_external.get(str(edge.get("target")))
        if not source or not target or source.id == target.id:
            continue
        relationship_type = "RELATED_TO"
        relationship = db.query(Relationship).filter_by(knowledge_base_id=document.knowledge_base_id, source_entity_id=source.id,
                                                         target_entity_id=target.id, relationship_type=relationship_type).first()
        props = edge.get("properties", {})
        description = str(props.get("description") or "")[:5000] or None
        if not relationship:
            relationship = Relationship(knowledge_base_id=document.knowledge_base_id, source_entity_id=source.id, target_entity_id=target.id,
                                        relationship_type=relationship_type, description=description, confidence=1.0,
                                        origin="ai_suggestion", review_status="suggested", is_legal=False)
            db.add(relationship); db.flush(); relationship_count += 1
            db.add(GraphProjectionEvent(event_type="relationship", relationship_id=relationship.id))
        elif relationship.deleted_at:
            relationship.deleted_at = None
        excerpt = description or f"{source.name} relates to {target.name}"
        if not db.query(RelationshipSource).filter_by(relationship_id=relationship.id, document_id=document.id, excerpt=excerpt).first():
            db.add(RelationshipSource(relationship_id=relationship.id, document_id=document.id, excerpt=excerpt)); relationship.source_count += 1
    db.commit()
    return {"entities": entity_count, "relationships": relationship_count}


LEGAL_GRAPH_RELATIONSHIPS = {
    "CONTAINS_PROVISION", "ISSUED_BY", "PARTY_TO", "REQUIRES", "GRANTS_RIGHT", "PROHIBITS", "DEFINES",
    "ISSUED_UNDER", "IMPLEMENTS", "AMENDS", "REPEALS", "SUPERSEDES", "REFERS_TO", "GOVERNED_BY",
}
# Cross-instrument subset of LEGAL_GRAPH_RELATIONSHIPS; the rest (CONTAINS_PROVISION,
# ISSUED_BY, PARTY_TO, ...) describe facts within a single document and never become
# a legal_instrument_relations row.
LEGAL_INSTRUMENT_RELATION_TYPES = {
    "ISSUED_UNDER", "IMPLEMENTS", "AMENDS", "REPEALS", "SUPERSEDES", "REFERS_TO", "GOVERNED_BY",
}


def _legal_value(item, *keys: str) -> str:
    if isinstance(item, str):
        return item.strip()
    if not isinstance(item, dict):
        return ""
    return str(next((item.get(key) for key in keys if item.get(key)), "")).strip()


def _legal_evidence(item, fallback: str) -> str:
    return (_legal_value(item, "evidence_quote", "excerpt", "text", "description") or fallback)[:5000]


def _reference_confidence(reference, default: float = 0.5) -> float:
    """An explicit confidence of 0.0 is a real (low) signal, not a missing value."""
    value = reference.get("confidence") if isinstance(reference, dict) else None
    return float(value) if value is not None else default


def legal_metadata_v2(document: Document) -> dict:
    """Normalize legacy legal metadata without discarding a human-edited value."""
    metadata = document.legal_metadata or {}
    if metadata.get("schema_version") == 2:
        # Backfill official corpus identity for records created before the
        # deterministic parser was introduced, while preserving curator edits.
        parsed = parse_legal_corpus_metadata(document.extracted_text or "", document.title or document.original_filename)
        current_instrument = metadata.get("instrument") if isinstance(metadata.get("instrument"), dict) else {}
        metadata["instrument"] = {**parsed.get("instrument", {}), **current_instrument}
        if not metadata.get("change_events") and parsed.get("change_events"):
            metadata["change_events"] = parsed["change_events"]
        if not metadata.get("amendments") and parsed.get("amendments"):
            metadata["amendments"] = parsed["amendments"]
        return metadata
    kind_by_document_type = {"legal": "Act", "regulation": "Regulation", "contract": "Contract"}
    provisions = []
    for article in metadata.get("articles", []) if isinstance(metadata.get("articles"), list) else []:
        number = _legal_value(article, "number", "article_number", "label", "name", "title")
        if number:
            provisions.append({
                "kind": _legal_value(article, "kind") or "article", "number": number,
                "heading": _legal_value(article, "heading", "title"), "text": _legal_value(article, "text"),
                "evidence_quote": _legal_evidence(article, f"มาตรา {number}"),
            })
    parsed = parse_legal_corpus_metadata(document.extracted_text or "", document.title or document.original_filename)
    parsed_instrument = parsed.get("instrument", {})
    return {
        "schema_version": 2,
        # Compatibility aliases remain readable by existing clients while the
        # graph projection uses the normalized v2 fields below.
        "document_type": metadata.get("document_type"),
        "articles": metadata.get("articles") if isinstance(metadata.get("articles"), list) else [],
        "instrument": {
            "kind": kind_by_document_type.get(document.document_type, "Other"),
            "official_title": document.title or document.original_filename,
            "official_number": metadata.get("document_number"),
            "jurisdiction": metadata.get("jurisdiction"),
            "effective_date": metadata.get("effective_date"),
            "issuer": metadata.get("issuer"),
            **parsed_instrument,
        },
        "provisions": provisions,
        "parties": metadata.get("parties") if isinstance(metadata.get("parties"), list) else [],
        "obligations": metadata.get("obligations") if isinstance(metadata.get("obligations"), list) else [],
        "rights": metadata.get("rights") if isinstance(metadata.get("rights"), list) else [],
        "prohibitions": metadata.get("prohibitions") if isinstance(metadata.get("prohibitions"), list) else [],
        "penalties": metadata.get("penalties") if isinstance(metadata.get("penalties"), list) else [],
        "definitions": metadata.get("definitions") if isinstance(metadata.get("definitions"), list) else [],
        "amendments": metadata.get("amendments") if isinstance(metadata.get("amendments"), list) else parsed.get("amendments", []),
        "change_events": metadata.get("change_events") if isinstance(metadata.get("change_events"), list) else parsed.get("change_events", []),
        "references": metadata.get("references") if isinstance(metadata.get("references"), list) else [],
        "confidence": metadata.get("confidence", 0.0),
        "provenance": {**parsed.get("provenance", {}), **(metadata.get("provenance") or {})},
    }


def _document_legal_identity(document: Document, suffix: str) -> str:
    return f"legal:{suffix}:{document.id}"


def _legal_fingerprint(value: str) -> str:
    return hashlib.sha256(canonical_entity_name(value).encode()).hexdigest()[:24]


def _remove_document_legal_projection(db: Session, document: Document) -> None:
    """Remove only generated evidence from one document; never touch manual rows."""
    relationship_ids = [row[0] for row in db.query(Relationship.id).join(RelationshipSource).filter(
        RelationshipSource.document_id == document.id, Relationship.origin == "legal_schema"
    ).all()]
    for relationship_id in relationship_ids:
        db.query(RelationshipSource).filter_by(relationship_id=relationship_id, document_id=document.id).delete(synchronize_session=False)
        relationship = db.get(Relationship, relationship_id)
        if relationship:
            relationship.source_count = db.query(func.count(RelationshipSource.id)).filter_by(relationship_id=relationship_id).scalar() or 0
            if relationship.source_count == 0:
                relationship.deleted_at = datetime.utcnow()
    entity_ids = [row[0] for row in db.query(Entity.id).join(EntitySource).filter(
        EntitySource.document_id == document.id, Entity.origin == "legal_schema"
    ).all()]
    for entity_id in entity_ids:
        db.query(EntitySource).filter_by(entity_id=entity_id, document_id=document.id).delete(synchronize_session=False)
        entity = db.get(Entity, entity_id)
        if entity:
            entity.source_count = db.query(func.count(EntitySource.id)).filter_by(entity_id=entity_id).scalar() or 0
            if entity.source_count == 0:
                entity.deleted_at = datetime.utcnow()
    db.flush()


def _upsert_legal_entity(db: Session, document: Document, *, name: str, entity_type: str, identity_key: str,
                         excerpt: str, attributes: dict | None = None, origin: str = "legal_schema",
                         review_status: str = "verified") -> tuple[Entity | None, bool]:
    name = str(name or "").strip()[:500]
    if not name:
        return None, False
    entity = db.query(Entity).filter_by(
        knowledge_base_id=document.knowledge_base_id, identity_key=identity_key, entity_type=entity_type,
    ).first()
    created = False
    if not entity:
        entity = Entity(
            knowledge_base_id=document.knowledge_base_id, name=name, canonical_name=canonical_entity_name(name),
            identity_key=identity_key, entity_type=entity_type, attributes=attributes or {}, confidence=1.0,
            origin=origin, review_status=review_status, is_legal=True,
        )
        db.add(entity); db.flush(); created = True
        db.add(GraphProjectionEvent(event_type="entity", entity_id=entity.id))
    else:
        entity.deleted_at, entity.name, entity.attributes = None, name, {**(entity.attributes or {}), **(attributes or {})}
        entity.origin, entity.review_status, entity.is_legal = origin, review_status, True
    excerpt = str(excerpt or name)[:5000]
    if not db.query(EntitySource).filter_by(entity_id=entity.id, document_id=document.id, excerpt=excerpt).first():
        db.add(EntitySource(entity_id=entity.id, document_id=document.id, excerpt=excerpt)); entity.source_count += 1
    return entity, created


def _upsert_legal_relationship(db: Session, document: Document, *, source: Entity | None, target: Entity | None,
                               relationship_type: str, excerpt: str, origin: str = "legal_schema",
                               review_status: str = "verified", confidence: float = 1.0, attributes: dict | None = None) -> tuple[Relationship | None, bool]:
    if not source or not target or source.id == target.id:
        return None, False
    relationship = db.query(Relationship).filter_by(
        knowledge_base_id=document.knowledge_base_id, source_entity_id=source.id,
        target_entity_id=target.id, relationship_type=relationship_type,
    ).first()
    created = False
    if not relationship:
        relationship = Relationship(
            knowledge_base_id=document.knowledge_base_id, source_entity_id=source.id, target_entity_id=target.id,
            relationship_type=relationship_type, confidence=confidence, attributes=attributes or {},
            origin=origin, review_status=review_status, is_legal=True,
        )
        db.add(relationship); db.flush(); created = True
        db.add(GraphProjectionEvent(event_type="relationship", relationship_id=relationship.id))
    else:
        # A reviewer decision is durable.  Rebuilding source metadata must not
        # silently turn an approved/rejected suggestion back into a pending one.
        if origin == "ai_suggestion" and relationship.origin == "ai_suggestion" and relationship.review_status in {"verified", "rejected"}:
            return relationship, False
        relationship.deleted_at, relationship.confidence = None, confidence
        relationship.attributes = {**(relationship.attributes or {}), **(attributes or {})}
        relationship.origin, relationship.review_status, relationship.is_legal = origin, review_status, True
    excerpt = str(excerpt or relationship_type)[:5000]
    if not db.query(RelationshipSource).filter_by(relationship_id=relationship.id, document_id=document.id, excerpt=excerpt).first():
        db.add(RelationshipSource(relationship_id=relationship.id, document_id=document.id, excerpt=excerpt)); relationship.source_count += 1
    return relationship, created


def legal_instrument_entity(db: Session, document: Document) -> Entity | None:
    return db.query(Entity).filter_by(
        knowledge_base_id=document.knowledge_base_id,
        identity_key=_document_legal_identity(document, "instrument"), entity_type="LegalInstrument",
    ).filter(Entity.deleted_at.is_(None)).first()


def sync_legal_document_graph(db: Session, document: Document, *, replace: bool = True) -> dict[str, int]:
    """Project evidence-backed Legal Graph v2 nodes and edges for one document."""
    if document.document_type not in LEGAL_DOCUMENT_TYPES:
        return {"entities": 0, "relationships": 0}
    if replace:
        _remove_document_legal_projection(db, document)
    if not document.legal_metadata:
        db.flush()
        return {"entities": 0, "relationships": 0}
    metadata = legal_metadata_v2(document)
    document.legal_metadata = metadata
    text = document.extracted_text or ""
    instrument = metadata.get("instrument") if isinstance(metadata.get("instrument"), dict) else {}
    title = _legal_value(instrument, "official_title", "title") or document.title or document.original_filename
    instrument_kind = _legal_value(instrument, "kind") or "Other"
    anchor, entity_created = _upsert_legal_entity(
        db, document, name=title, entity_type="LegalInstrument", identity_key=_document_legal_identity(document, "instrument"),
        excerpt=text[:500] or title, attributes={"instrument_kind": instrument_kind, "document_id": document.id,
        "official_number": instrument.get("official_number"), "effective_date": instrument.get("effective_date")},
    )
    entity_count, relationship_count = int(entity_created), 0

    issuer = instrument.get("issuer")
    issuer_name = _legal_value(issuer, "name", "organization")
    if issuer_name:
        issuer_entity, created = _upsert_legal_entity(
            db, document, name=issuer_name, entity_type="LegalAuthority",
            identity_key=f"legal:authority:{canonical_entity_name(issuer_name)}", excerpt=_legal_evidence(issuer, issuer_name),
        ); entity_count += int(created)
        _, created = _upsert_legal_relationship(db, document, source=anchor, target=issuer_entity, relationship_type="ISSUED_BY", excerpt=_legal_evidence(issuer, issuer_name)); relationship_count += int(created)

    for provision in metadata.get("provisions", []) if isinstance(metadata.get("provisions"), list) else []:
        number = _legal_value(provision, "number", "label")
        if not number:
            continue
        kind = _legal_value(provision, "kind") or "article"
        label = f"{kind.title()} {number}" if not str(number).startswith(("มาตรา", "ข้อ", "Article", "Clause")) else str(number)
        provision_entity, created = _upsert_legal_entity(
            db, document, name=label, entity_type="Provision",
            identity_key=f"legal:provision:{document.id}:{canonical_entity_name(kind)}:{canonical_entity_name(number)}",
            excerpt=_legal_evidence(provision, label), attributes={"provision_kind": kind, "provision_number": str(number), "heading": _legal_value(provision, "heading")},
        ); entity_count += int(created)
        _, created = _upsert_legal_relationship(db, document, source=anchor, target=provision_entity, relationship_type="CONTAINS_PROVISION", excerpt=_legal_evidence(provision, label)); relationship_count += int(created)

    for values_key, entity_type, relationship_type in [
        ("parties", "LegalParty", "PARTY_TO"), ("obligations", "Obligation", "REQUIRES"),
        ("rights", "Right", "GRANTS_RIGHT"), ("prohibitions", "Prohibition", "PROHIBITS"),
        ("definitions", "Definition", "DEFINES"), ("penalties", "Penalty", "REQUIRES"),
        ("amendments", "Amendment", "AMENDS"),
    ]:
        values = metadata.get(values_key) if isinstance(metadata.get(values_key), list) else []
        for item in values:
            name = _legal_value(item, "name", "party", "organization", "description", "obligation", "title", "term", "definition", "penalty")
            if not name:
                continue
            key = f"legal:{entity_type.casefold()}:{document.id}:{_legal_fingerprint(name)}"
            item_entity, created = _upsert_legal_entity(db, document, name=name, entity_type=entity_type, identity_key=key,
                excerpt=_legal_evidence(item, name), attributes={"role": _legal_value(item, "role")})
            entity_count += int(created)
            _, created = _upsert_legal_relationship(db, document, source=anchor, target=item_entity,
                relationship_type=relationship_type, excerpt=_legal_evidence(item, name))
            relationship_count += int(created)
    db.flush()
    return {"entities": entity_count, "relationships": relationship_count}


def _get_or_create_legal_family(db: Session, knowledge_base_id: str, title: str, normalized_key: str) -> LegalFamily:
    family = db.query(LegalFamily).filter_by(knowledge_base_id=knowledge_base_id, normalized_key=normalized_key).first()
    if family is None:
        family = LegalFamily(knowledge_base_id=knowledge_base_id, base_title=title[:500], normalized_key=normalized_key)
        db.add(family); db.flush()
    return family


def upsert_legal_instrument(db: Session, document: Document) -> LegalInstrument | None:
    """Register or refresh the legal registry entry for one legal document.

    Never touches `status`/`status_reason`; only resolve_instrument_statuses does,
    so extraction re-runs cannot clobber a status derived from reviewed relations.
    """
    if document.document_type not in LEGAL_DOCUMENT_TYPES or not document.legal_metadata:
        return None
    metadata = legal_metadata_v2(document)
    instrument_meta = metadata.get("instrument") if isinstance(metadata.get("instrument"), dict) else {}
    title = _legal_value(instrument_meta, "official_title", "title") or document.title or document.original_filename
    row = db.query(LegalInstrument).filter_by(document_id=document.id).first()
    if row is None:
        row = LegalInstrument(document_id=document.id, knowledge_base_id=document.knowledge_base_id,
                              status_source="resolver", review_status="unreviewed")
        db.add(row)
    if row.status_source != "manual":
        base_title_key, version_label, enacted_year = normalize_family_key(title)
        legal_work_key = _legal_value(instrument_meta, "legal_work_key") or base_title_key
        # Header parser canonicalizes the common Thai typo and amendment prefix.
        legal_work_key = canonical_entity_name(legal_work_key)
        family = _get_or_create_legal_family(db, document.knowledge_base_id, title, legal_work_key)
        row.family_id = family.id
        row.kind = classify_kind(title, _legal_value(instrument_meta, "kind") or None)
        row.authority_level = AUTHORITY_LEVELS.get(row.kind, AUTHORITY_LEVELS["other"])
        row.official_title = title[:500]
        row.official_number = _legal_value(instrument_meta, "official_number") or None
        row.issuer = _legal_value(instrument_meta.get("issuer"), "name", "organization") or None
        row.jurisdiction = _legal_value(instrument_meta, "jurisdiction") or None
        row.version_label = version_label
        row.enacted_year = enacted_year
        row.legal_work_key = legal_work_key
        row.document_class = _legal_value(instrument_meta, "document_class") or None
        row.version_role = _legal_value(instrument_meta, "version_role") or row.document_class or None
        row.version_date = parse_thai_date(instrument_meta.get("version_date")) or document.published_at
        row.effective_from = parse_thai_date(instrument_meta.get("effective_date")) or parse_thai_date(instrument_meta.get("version_date")) or document.published_at
        row.effective_to = parse_thai_date(instrument_meta.get("effective_to"))
        # Extraction may provide an official source, but a curator's manual
        # provenance must never be overwritten by a re-process.
        if not row.source_uri:
            row.source_uri = _legal_value(instrument_meta, "source_uri", "official_source_url", "source_url") or None
        if not row.source_reference:
            row.source_reference = _legal_value(instrument_meta, "source_reference", "gazette_reference", "ราชกิจจานุเบกษา") or None
    db.flush()
    return row


def _sync_deterministic_change_events(db: Session, knowledge_base_id: str) -> int:
    """Materialize explicit amendment/repeal clauses as verified registry edges.

    The target is the latest consolidated expression at or before the amendment
    date.  This prevents a 1977 amendment from being linked to the 2019 text.
    """
    instruments = db.query(LegalInstrument).filter_by(knowledge_base_id=knowledge_base_id).all()
    docs = {doc.id: doc for doc in db.query(Document).filter(Document.id.in_([row.document_id for row in instruments])).all()}
    changed = 0
    # Two clauses in one act can normalize to the same edge (e.g. a "replace" and
    # an "insert" that both touch มาตรา 105 both map to AMENDS→105). The find-or-
    # create query cannot see a sibling added earlier in this same run, so guard
    # the unique key in-memory to avoid a duplicate-key flush failure.
    batch_seen: set[tuple[str, str, str, str | None]] = set()
    for source in instruments:
        if source.document_class != "amendment":
            continue
        source_doc = docs.get(source.document_id)
        events = (legal_metadata_v2(source_doc).get("change_events", []) if source_doc else [])
        if not events:
            continue
        candidates = [row for row in instruments if row.id != source.id and row.legal_work_key == source.legal_work_key and row.document_class in {"main", "consolidated"}]
        source_date = source.version_date or source.effective_from
        prior = [row for row in candidates if not source_date or not row.version_date or row.version_date <= source_date]
        target = max(prior or candidates, key=lambda row: row.version_date or row.effective_from or datetime.min.date(), default=None)
        if not target:
            continue
        for event in events:
            action = str(event.get("action") or "replace").casefold()
            relation_type = "REPEALS" if action == "repeal" else "AMENDS"
            target_provision = str(event.get("target_provision") or event.get("provision_number") or "")[:120] or None
            batch_key = (source.id, relation_type, target.id, target_provision)
            if batch_key in batch_seen:
                continue
            batch_seen.add(batch_key)
            row = db.query(LegalInstrumentRelation).filter_by(
                source_instrument_id=source.id, relation=relation_type,
                target_instrument_id=target.id, target_provision=target_provision,
            ).first()
            if not row:
                row = LegalInstrumentRelation(knowledge_base_id=knowledge_base_id,
                    source_instrument_id=source.id, target_instrument_id=target.id,
                    relation=relation_type, target_provision=target_provision,
                    origin="legal_schema", review_status="verified")
                db.add(row); changed += 1
            row.evidence_quote = str(event.get("evidence_quote") or "")[:5000]
            row.confidence, row.origin, row.review_status = 1.0, "legal_schema", "verified"
    db.flush()
    return changed


def _sync_legal_instrument_relation(db: Session, document: Document, target_document: Document | None,
                                    relationship: Relationship | None, relationship_type: str, reference: dict) -> None:
    """Mirror one cross-document suggestion into the legal registry so the status
    resolver and query-time resolver can reason about it without re-parsing text."""
    if relationship_type not in LEGAL_INSTRUMENT_RELATION_TYPES:
        return
    source_instrument = db.query(LegalInstrument).filter_by(document_id=document.id).first()
    if not source_instrument:
        return
    target_instrument = db.query(LegalInstrument).filter_by(document_id=target_document.id).first() if target_document else None
    target_provision = _legal_value(reference, "target_provision") or None
    row = db.query(LegalInstrumentRelation).filter_by(
        source_instrument_id=source_instrument.id, relation=relationship_type,
        target_instrument_id=target_instrument.id if target_instrument else None,
        target_provision=target_provision,
    ).first()
    if row and row.origin == "manual":
        return
    if not row:
        row = LegalInstrumentRelation(
            knowledge_base_id=document.knowledge_base_id, source_instrument_id=source_instrument.id,
            relation=relationship_type, origin="ai_suggestion", review_status="suggested",
        )
        db.add(row)
    row.target_instrument_id = target_instrument.id if target_instrument else None
    row.target_text = (_legal_value(reference, "target_title", "title", "instrument_title") or None or "")[:700] or None
    row.target_provision = target_provision
    row.evidence_quote = _legal_evidence(reference, relationship_type)[:5000]
    row.confidence = _reference_confidence(reference)
    row.relationship_id = relationship.id if relationship else None
    if relationship and relationship.review_status in {"verified", "rejected"}:
        row.review_status = relationship.review_status
    db.flush()


def sync_legal_instrument_relation_review(db: Session, relationship: Relationship) -> None:
    """Propagate an admin's approve/reject decision to the registry and re-resolve status."""
    rows = db.query(LegalInstrumentRelation).filter_by(relationship_id=relationship.id).all()
    manual_rows = [row for row in rows if row.origin != "manual"]
    for row in manual_rows:
        row.review_status = relationship.review_status
    db.flush()
    for knowledge_base_id in {row.knowledge_base_id for row in manual_rows}:
        resolve_instrument_statuses(db, knowledge_base_id)


def _clear_legal_suggestions(db: Session, knowledge_base_id: str) -> None:
    rows = db.query(Relationship).filter_by(knowledge_base_id=knowledge_base_id, origin="ai_suggestion", is_legal=True, review_status="suggested").all()
    for relationship in rows:
        relationship.deleted_at = datetime.utcnow()


def _target_instrument(db: Session, document: Document, reference: dict) -> Document | None:
    title = _legal_value(reference, "target_title", "title", "instrument_title")
    number = _legal_value(reference, "target_number", "official_number", "instrument_number")
    candidates = db.query(Document).filter(
        Document.knowledge_base_id == document.knowledge_base_id, Document.id != document.id,
        Document.document_type.in_(LEGAL_DOCUMENT_TYPES), Document.status == "completed", Document.deleted_at.is_(None),
    ).all()
    matches = []
    source_meta = legal_metadata_v2(document).get("instrument") or {}
    source_work = canonical_entity_name(_legal_value(source_meta, "legal_work_key"))
    source_date = parse_thai_date(source_meta.get("version_date")) or document.published_at
    for candidate in candidates:
        metadata = legal_metadata_v2(candidate)
        instrument = metadata.get("instrument") or {}
        candidate_title = _legal_value(instrument, "official_title") or candidate.title or candidate.original_filename
        candidate_work = canonical_entity_name(_legal_value(instrument, "legal_work_key"))
        if source_work and candidate_work and source_work != candidate_work:
            continue
        score = 0
        if number and number == str(instrument.get("official_number") or ""):
            score += 100
        if title and (canonical_entity_name(title) in canonical_entity_name(candidate_title) or canonical_entity_name(candidate_title) in canonical_entity_name(title)):
            score += 50
        if score:
            candidate_date = parse_thai_date(instrument.get("version_date")) or candidate.published_at
            # Prefer the latest expression not newer than the referring act.
            if source_date and candidate_date and candidate_date > source_date:
                score -= 20
            matches.append((score, candidate_date or datetime.min.date(), candidate))
    return max(matches, key=lambda item: (item[0], item[1]))[2] if matches else None


def build_legal_cross_document_suggestions(db: Session, knowledge_base_id: str, *, allow_model: bool = True) -> int:
    """Create reviewable, evidenced cross-instrument links from extracted references."""
    count = 0
    documents = db.query(Document).filter(
        Document.knowledge_base_id == knowledge_base_id, Document.document_type.in_(LEGAL_DOCUMENT_TYPES),
        Document.status == "completed", Document.deleted_at.is_(None),
    ).all()
    for document in documents:
        source = legal_instrument_entity(db, document)
        metadata = legal_metadata_v2(document)
        references = metadata.get("references", []) if isinstance(metadata.get("references"), list) else []
        if allow_model and not references and document.extracted_text:
            candidates = []
            for candidate in documents:
                if candidate.id == document.id:
                    continue
                candidate_metadata = legal_metadata_v2(candidate)
                candidate_instrument = candidate_metadata.get("instrument") or {}
                candidates.append({
                    "title": _legal_value(candidate_instrument, "official_title") or candidate.title or candidate.original_filename,
                    "official_number": str(candidate_instrument.get("official_number") or ""),
                })
            references = OpenRouterClient().suggest_legal_relationships(document.title or document.original_filename, document.extracted_text, candidates)
            if references:
                metadata["references"] = references
                document.legal_metadata = metadata
        for reference in references:
            if not isinstance(reference, dict):
                continue
            relationship_type = _legal_value(reference, "relationship", "relationship_type").upper() or "REFERS_TO"
            if relationship_type not in LEGAL_GRAPH_RELATIONSHIPS:
                relationship_type = "REFERS_TO"
            target_document = _target_instrument(db, document, reference)
            target = legal_instrument_entity(db, target_document) if target_document else None
            relationship, created = _upsert_legal_relationship(
                db, document, source=source, target=target, relationship_type=relationship_type,
                excerpt=_legal_evidence(reference, relationship_type), origin="ai_suggestion", review_status="suggested",
                confidence=_reference_confidence(reference), attributes={"reference": reference},
            )
            count += int(created)
            _sync_legal_instrument_relation(db, document, target_document, relationship, relationship_type, reference)
    db.flush()
    return count


def rebuild_legal_graph(db: Session, knowledge_base_id: str) -> dict[str, int]:
    documents = db.query(Document).filter(
        Document.knowledge_base_id == knowledge_base_id, Document.document_type.in_(LEGAL_DOCUMENT_TYPES),
        Document.status == "completed", Document.deleted_at.is_(None),
    ).all()
    totals = {"documents": len(documents), "entities": 0, "relationships": 0, "suggestions": 0}
    _clear_legal_suggestions(db, knowledge_base_id)
    for document in documents:
        result = sync_legal_document_graph(db, document, replace=True)
        totals["entities"] += result["entities"]; totals["relationships"] += result["relationships"]
        upsert_legal_instrument(db, document)
        link_provisions_to_chunks(db, document)
    totals["deterministic_relations"] = _sync_deterministic_change_events(db, knowledge_base_id)
    # Registry rebuild is a deterministic maintenance operation.  It must not
    # fan out into one model call per document; curated/previously-extracted
    # references are still materialized, while optional AI suggestions can be
    # generated explicitly through the existing helper when desired.
    totals["suggestions"] = build_legal_cross_document_suggestions(db, knowledge_base_id, allow_model=False)
    totals["status_changes"] = resolve_instrument_statuses(db, knowledge_base_id)["changed"]
    db.commit()
    return totals


def resolve_entity(db: Session, knowledge_base_ids: list[str], text: str) -> Entity | None:
    canonical = canonical_entity_name(text)
    query = db.query(Entity).filter(Entity.deleted_at.is_(None))
    if knowledge_base_ids:
        query = query.filter(Entity.knowledge_base_id.in_(knowledge_base_ids))
    exact = query.filter(Entity.canonical_name == canonical).first()
    if exact:
        return exact
    candidates = query.filter(Entity.name.ilike(f"%{text}%")).all()
    if candidates:
        return candidates[0]
    # JSON aliases are portable between PostgreSQL and SQLite when matched in
    # Python; the graph is bounded and this runs only after indexed lookups.
    for candidate in query.limit(500).all():
        if canonical in {canonical_entity_name(alias) for alias in (candidate.aliases or [])}:
            return candidate
    return None


def relationship_sources(db: Session, relationships: list[Relationship], plan: RetrievalPlan | None = None) -> list[dict]:
    ids = [relationship.id for relationship in relationships]
    if not ids:
        return []
    rows = db.query(RelationshipSource, Document).join(Document, Document.id == RelationshipSource.document_id).filter(
        RelationshipSource.relationship_id.in_(ids), Document.status == "completed", Document.deleted_at.is_(None)
    )
    if plan and plan.published_from:
        rows = rows.filter(Document.published_at >= plan.published_from)
    if plan and plan.published_to:
        rows = rows.filter(Document.published_at < plan.published_to)
    rows = _apply_metadata_filter(rows, plan)
    rows = rows.all()
    return [{"citation_id": f"S{index}", "document_id": document.id, "title": document.title,
             "chunk_id": source.id, "excerpt": source.excerpt, "relevance": 1.0 / index,
             "_relationship_id": source.relationship_id}
            for index, (source, document) in enumerate(rows, 1)]


def _relationship_filter_active(plan: RetrievalPlan | None) -> bool:
    return bool(plan and (plan.published_from or plan.published_to or plan.metadata_document_ids is not None))


def _filter_relationships_by_plan(db: Session, relationships: list[Relationship], plan: RetrievalPlan | None) -> list[Relationship]:
    """Keep only edges with at least one source document in the hard-filter scope."""
    if not relationships or not _relationship_filter_active(plan):
        return relationships
    query = db.query(RelationshipSource.relationship_id).join(
        Document, Document.id == RelationshipSource.document_id,
    ).filter(
        RelationshipSource.relationship_id.in_([relationship.id for relationship in relationships]),
        Document.status == "completed", Document.deleted_at.is_(None),
    )
    query = _apply_published_filter(query, plan)
    query = _apply_metadata_filter(query, plan)
    allowed = {row[0] for row in query.distinct().all()}
    return [relationship for relationship in relationships if relationship.id in allowed]


def entity_graph(db: Session, entity: Entity, depth: int = 1) -> dict:
    visited = {entity.id}
    frontier = {entity.id}
    edges: list[Relationship] = []
    for _ in range(depth):
        found = db.query(Relationship).filter(Relationship.deleted_at.is_(None)).filter(
            (Relationship.is_legal.is_(False)) | (Relationship.review_status == "verified")
        ).filter(Relationship.knowledge_base_id == entity.knowledge_base_id).filter(
            (Relationship.source_entity_id.in_(frontier)) | (Relationship.target_entity_id.in_(frontier))).all()
        edges.extend(edge for edge in found if edge.id not in {item.id for item in edges})
        next_frontier = {edge.source_entity_id for edge in found} | {edge.target_entity_id for edge in found}
        next_frontier -= visited
        visited |= next_frontier
        frontier = next_frontier
        if not frontier:
            break
    nodes = db.query(Entity).filter(Entity.id.in_(visited)).all()
    return {
        "nodes": [{"id": node.id, "name": node.name, "type": node.entity_type} for node in nodes],
        "edges": [{"id": edge.id, "source": edge.source_entity_id, "target": edge.target_entity_id, "type": edge.relationship_type} for edge in edges],
    }


def analyze_impact(db: Session, subject: str, knowledge_base_ids: list[str], max_depth: int, include_indirect: bool, entity_id: str | None = None) -> dict:
    entity = db.get(Entity, entity_id) if entity_id else resolve_entity(db, knowledge_base_ids, subject)
    if entity and knowledge_base_ids and entity.knowledge_base_id not in knowledge_base_ids:
        entity = None
    if not entity:
        return {"status": "success", "subject": None, "direct_impacts": [], "indirect_impacts": [], "sources": [], "insufficient_evidence": True, "warnings": ["Entity was not found."]}
    graph = entity_graph(db, entity, max_depth if include_indirect else 1)
    nodes = {node["id"]: node for node in graph["nodes"]}
    edges_by_node: dict[str, list[dict]] = {}
    for edge in graph["edges"]:
        edges_by_node.setdefault(edge["source"], []).append(edge)
        edges_by_node.setdefault(edge["target"], []).append(edge)
    distance = {entity.id: 0}
    paths = {entity.id: [entity.id]}
    frontier = [entity.id]
    selected_edges: list[Relationship] = []
    while frontier:
        current = frontier.pop(0)
        if distance[current] >= max_depth:
            continue
        for edge in edges_by_node.get(current, []):
            other = edge["target"] if edge["source"] == current else edge["source"]
            if other in distance:
                continue
            distance[other] = distance[current] + 1
            paths[other] = paths[current] + [other]
            frontier.append(other)
            selected_edges.append(db.get(Relationship, edge["id"]))
    sources = relationship_sources(db, [edge for edge in selected_edges if edge])
    citation_ids_by_rel: dict[str, list[str]] = {}
    for source in sources:
        citation_ids_by_rel.setdefault(source.pop("_relationship_id"), []).append(source["citation_id"])
    direct, indirect = [], []
    for node_id, node in nodes.items():
        if node_id == entity.id or node_id not in distance:
            continue
        edge = next((item for item in selected_edges if item and (item.source_entity_id == node_id or item.target_entity_id == node_id)), None)
        item = {"entity_id": node_id, "name": node["name"], "type": node["type"], "distance": distance[node_id],
                "path": [nodes[path_id]["name"] for path_id in paths[node_id]], "relationship": edge.relationship_type if edge else "RELATED_TO",
                "confidence": edge.confidence if edge else None, "citation_ids": citation_ids_by_rel.get(edge.id, []) if edge else []}
        (direct if distance[node_id] == 1 else indirect).append(item)
    return {"status": "success", "subject": {"entity_id": entity.id, "name": entity.name, "type": entity.entity_type},
            "direct_impacts": direct, "indirect_impacts": indirect if include_indirect else [], "sources": sources,
            "insufficient_evidence": not bool(direct or indirect), "warnings": []}


def store_upload(upload, knowledge_base_id: str) -> tuple[str, str, int, str, str]:
    settings = get_settings()
    extension = Path(upload.filename or "").suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("FILE_TYPE_NOT_SUPPORTED")
    reported_mime = (upload.content_type or "").lower()
    if reported_mime and reported_mime not in {"application/octet-stream", "binary/octet-stream"} and reported_mime not in ALLOWED_MIME_TYPES[extension]:
        raise ValueError("FILE_MIME_TYPE_NOT_SUPPORTED")
    root = settings.file_root / knowledge_base_id
    root.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid.uuid4()}{extension}"
    destination = root / stored_name
    size = 0
    digest = hashlib.sha256()
    with destination.open("wb") as output:
        while chunk := upload.file.read(1024 * 1024):
            size += len(chunk)
            if size > settings.max_file_size_mb * 1024 * 1024:
                destination.unlink(missing_ok=True)
                raise ValueError("FILE_TOO_LARGE")
            digest.update(chunk)
            output.write(chunk)
    mime_type = upload.content_type or mimetypes.guess_type(upload.filename or "")[0] or "application/octet-stream"
    return str(destination), stored_name, size, digest.hexdigest(), mime_type


DOCUMENT_TYPES = {"general", "legal", "regulation", "contract"}
LEGAL_DOCUMENT_TYPES = {"legal", "regulation", "contract"}

DOCUMENT_TYPE_LABELS = {
    "general": "General document",
    "legal": "Legal document",
    "regulation": "Regulation / policy",
    "contract": "Contract",
}
LEGAL_KIND_LABELS = {
    "constitution": "รัฐธรรมนูญ",
    "act": "พระราชบัญญัติ",
    "royal_decree": "พระราชกฤษฎีกา",
    "ministerial_regulation": "กฎกระทรวง",
    "notification": "ประกาศ",
    "rule": "ระเบียบ/ข้อบังคับ",
    "circular": "หนังสือเวียน",
    "guideline": "แนวปฏิบัติ/คู่มือ",
    "resolution": "มติ",
    "other": "อื่น ๆ",
}
DOCUMENT_CLASS_LABELS = {
    "main": "กฎหมายหลัก",
    "consolidated": "ฉบับรวม/ฉบับปรับปรุง",
    "amendment": "ฉบับแก้ไขเพิ่มเติม",
}

VERSION_ROLE_LABELS = {
    "main": "กฎหมายหลัก",
    "amendment": "ฉบับแก้ไขเพิ่มเติม",
    "consolidated": "ฉบับรวม/ฉบับปรับปรุง",
    "latest_consolidated": "ฉบับปรับปรุงล่าสุด",
    "unknown": "ไม่ระบุชั้นฉบับ",
}

_INVENTORY_COUNT_TERMS = ("กี่ฉบับ", "กี่เอกสาร", "จำนวน", "รายการ", "รายชื่อ", "แสดงรายการ", "มีอะไรบ้าง", "how many", "count", "list", "show", "what documents")
_INVENTORY_TYPE_TERMS = ("ประเภท", "แบ่งกลุ่ม", "แบ่งเป็น", "category", "categories", "grouped")
_INVENTORY_SUBJECT_TERMS = ("เอกสาร", "กฎหมาย", "knowledge base", "document", "law")
_INVENTORY_CURRENT_TERMS = ("ปัจจุบัน", "ล่าสุด", "มีผลบังคับใช้", "ยังมีผล", "current", "latest", "in force")
_INVENTORY_ACTION_TERMS = ("ลบ", "แก้ไข", "อัปโหลด", "อัพโหลด", "ดาวน์โหลด", "ย้าย", "เก็บ", "archive", "delete", "remove", "upload", "download", "move", "process", "reindex", "how do i", "อย่างไร", "ทำอย่างไร", "วิธี")
_LEGAL_METADATA_LOOKUP_TERMS = ("วันที่ในผัง", "ฉบับล่าสุดในชุด", "เอกสารใด", "ต้องคืนเอกสาร", "ข้อความจากกฎหมายฉบับ")
_LEGAL_PROVENANCE_TERMS = ("แก้ไขโดย", "แก้ไขเพิ่มเติมโดย", "ถูกเพิ่มโดย", "เพิ่มโดย", "ยกเลิกโดย", "ฉบับใด")
_THAI_DIGIT_TRANSLATION = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")


def is_document_inventory_query(query: str) -> bool:
    """Recognize count/list/type questions before the LLM planner runs."""
    value = (query or "").strip().casefold()
    if not value or not any(term in value for term in _INVENTORY_SUBJECT_TERMS):
        return False
    count_signal = any(term in value for term in _INVENTORY_COUNT_TERMS)
    type_signal = any(term in value for term in _INVENTORY_TYPE_TERMS)
    if any(term in value for term in _INVENTORY_ACTION_TERMS) and not any(term in value for term in ("กี่", "จำนวน", "how many", "count")):
        return False
    return count_signal or (type_signal and "แบ่ง" in value)


def document_inventory_scope(query: str, query_filters=None) -> str:
    if bool(getattr(query_filters, "include_historical", False)):
        return "all"
    value = (query or "").casefold()
    return "current" if any(term in value for term in _INVENTORY_CURRENT_TERMS) else "all"


def _requested_amendment_number(query: str) -> str | None:
    match = re.search(r"ฉบับ(?:แก้ไข)?\s*(?:ครั้ง)?\s*ที่\s*([0-9๐-๙]+)", query or "", re.IGNORECASE)
    return match.group(1).translate(_THAI_DIGIT_TRANSLATION) if match else None


def is_legal_metadata_lookup(query: str) -> bool:
    """Recognize document/version questions that must not use chunk similarity.

    This is intentionally based on reusable legal metadata language, not a
    corpus name or a document filename.  Questions asking how an amendment
    changes a provision continue through the provision/provenance path.
    """
    value = (query or "").casefold()
    asks_latest = "ฉบับล่าสุด" in value and any(term in value for term in ("เอกสาร", "ชุดข้อมูล", "วันที่ในผัง", "ฉบับใด"))
    asks_exact = bool(_requested_amendment_number(query)) and any(term in value for term in _LEGAL_METADATA_LOOKUP_TERMS)
    return asks_latest or asks_exact


def _legal_metadata_source(document: Document, instrument: LegalInstrument) -> dict:
    title = document.title or document.original_filename
    return {
        "citation_id": "S1", "document_id": document.id, "title": title,
        "excerpt": f"{instrument.official_title or title}; วันที่ในผัง {instrument.version_date.isoformat() if instrument.version_date else 'ไม่ระบุ'}",
        "relevance": 1.0, "document_status": instrument.status,
        "document_class": instrument.document_class, "version_role": instrument.version_role,
        "version_label": instrument.version_label, "version_date": instrument.version_date.isoformat() if instrument.version_date else None,
        "source_uri": instrument.source_uri, "source_reference": instrument.source_reference,
        "legal_label": instrument.official_title or title,
    }


def _trace_preview_fields(query: str, answer: str, sources: list[dict]) -> dict:
    """Populate the same observability fields the generic retrieval path and the
    document-inventory shortcut record (query_preview/answer_preview/etc.).

    Every deterministic legal-registry shortcut builds its own result dict
    rather than going through the generic composer, so without this they show
    "Query preview unavailable" / "No answer preview" in the Trace Explorer even
    though the query was answered successfully.
    """
    settings = get_settings()
    query = query or ""
    return {
        "query_preview": query[:500] if settings.log_query_text else None,
        "query_length": len(query),
        "query_sha256": hashlib.sha256(query.encode("utf-8")).hexdigest(),
        "answer_preview": (answer or "")[:800] if settings.log_query_text else None,
        "citation_ids": [source.get("citation_id") for source in sources],
        "filter_summary": {},
    }


def build_legal_metadata_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    """Resolve publisher/version metadata directly from the legal registry."""
    rows = db.query(Document, LegalInstrument).join(LegalInstrument, LegalInstrument.document_id == Document.id).filter(
        Document.knowledge_base_id.in_(kb_ids), Document.deleted_at.is_(None), Document.status == "completed"
    )
    requested = _requested_amendment_number(query)
    if requested:
        rows = rows.filter(LegalInstrument.version_role == "amendment", LegalInstrument.official_number == requested)
    else:
        rows = rows.filter(LegalInstrument.version_role == "latest_consolidated")
    matches = rows.order_by(LegalInstrument.version_date.desc(), Document.created_at.desc()).all()
    sources = [_legal_metadata_source(document, instrument) for document, instrument in matches[:1]]
    if sources:
        document, instrument = matches[0]
        date_text = instrument.version_date.isoformat() if instrument.version_date else "ไม่ระบุวันที่ในผัง"
        if requested:
            answer = f"เอกสารที่ตรงกับฉบับแก้ไขครั้งที่ {requested} คือ {document.title or document.original_filename} — {instrument.official_title or ''} — วันที่ในผัง {date_text} [S1]"
        else:
            answer = f"ฉบับล่าสุดในชุดข้อมูลคือ {document.title or document.original_filename} — {instrument.official_title or ''} — วันที่ในผัง {date_text} [S1]"
        answer += "\n\nรายละเอียดแหล่งอ้างอิง:\n" + _render_citation_details(sources)
    else:
        answer = "ไม่พบเอกสารกฎหมายที่มี metadata ตรงกับฉบับที่ร้องขอในขอบเขตของ MCP key"
    result = {
        "status": "success", "result_id": "", "answer": answer, "insufficient_evidence": not bool(sources),
        "sources": sources, "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": "legal_metadata_registry",
        "retrieval_plan": {"version": 1, "intent": "legal_metadata_lookup", "planner_source": "rules", "channels": ["legal_registry"]},
        "retrieval_trace": [{"channel": "legal_registry", "system": "PostgreSQL legal registry", "status": "used", "result_count": len(sources), "duration_ms": 0, "detail": "publisher metadata lookup"}],
        **_trace_preview_fields(query, answer, sources)},
    }
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result


def is_legal_provenance_lookup(query: str) -> bool:
    """Detect amendment/repeal questions answerable from verified registry edges."""
    return bool(parse_provision_refs(query)) and any(term in (query or "").casefold() for term in _LEGAL_PROVENANCE_TERMS)


def is_legal_effective_rule_lookup(query: str) -> bool:
    value = (query or "").casefold()
    return bool(_requested_amendment_number(query)) and any(term in value for term in ("มีผลใช้บังคับ", "เริ่มใช้บังคับ", "กำหนดเวลา", "วันถัดจากวันประกาศ", "วันประกาศ"))


def build_legal_effective_rule_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    """Return the operative commencement/change clauses of a named amendment.

    This retrieves the statute's own sections 2--3 instead of treating the
    publisher's version date as a legal effective date.
    """
    requested = _requested_amendment_number(query)
    rows = db.query(Document, LegalInstrument).join(LegalInstrument, LegalInstrument.document_id == Document.id).filter(
        Document.knowledge_base_id.in_(kb_ids), Document.deleted_at.is_(None), Document.status == "completed",
        LegalInstrument.version_role == "amendment", LegalInstrument.official_number == requested,
    ).order_by(LegalInstrument.version_date.desc()).all()
    sources: list[dict] = []
    if rows:
        document, instrument = rows[0]
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document.id, DocumentChunk.section_number.in_(("2", "3")),
        ).order_by(DocumentChunk.chunk_index).all()
        for chunk in chunks:
            sources.append({
                "citation_id": f"S{len(sources) + 1}", "document_id": document.id, "chunk_id": chunk.id,
                "title": document.title or document.original_filename, "section_label": chunk.section_label,
                "excerpt": chunk.content, "relevance": 1.0, "document_status": instrument.status,
                "document_class": instrument.document_class, "version_role": instrument.version_role,
                "source_uri": instrument.source_uri, "source_reference": instrument.source_reference,
                "legal_label": instrument.official_title or document.title or document.original_filename,
            })
    if sources:
        answer = "\n\n".join(f"{source['excerpt']}\n[{source['citation_id']}]" for source in sources)
        answer += "\n\nรายละเอียดแหล่งอ้างอิง:\n" + _render_citation_details(sources)
    else:
        answer = "ไม่พบข้อบทว่าด้วยวันเริ่มใช้บังคับของฉบับแก้ไขที่ระบุในคลังข้อมูลที่เลือก"
    result = {
        "status": "success", "result_id": "", "answer": answer, "insufficient_evidence": not bool(sources), "sources": sources,
        "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": "legal_effective_rule_registry",
        "retrieval_plan": {"version": 1, "intent": "legal_effective_date_rule", "planner_source": "rules", "channels": ["legal_registry", "document_chunks"]},
        "retrieval_trace": [{"channel": "legal_registry", "system": "PostgreSQL legal registry", "status": "used", "result_count": len(sources), "duration_ms": 0, "detail": "named amendment commencement clauses"}],
        **_trace_preview_fields(query, answer, sources)},
    }
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result


def _verified_amendment_relations(db: Session, kb_ids: list[str], family_id: str | None) -> list[tuple[LegalInstrumentRelation, LegalInstrument, Document]]:
    """Fetch every verified AMENDS relation once, scoped to one legal family.

    Scoping by the amended (target) instrument's family_id keeps amendment
    attribution from being answered by a different, unrelated law that
    happens to share the same bare section number (e.g. two unrelated codes
    both loaded into one Knowledge Base each having their own "มาตรา 104").
    Fetching once per caller instead of once per provision also avoids
    repeating this full-table scan when a result names several sections.
    """
    query = db.query(LegalInstrumentRelation, LegalInstrument, Document).join(
        LegalInstrument, LegalInstrument.id == LegalInstrumentRelation.source_instrument_id
    ).join(Document, Document.id == LegalInstrument.document_id).filter(
        LegalInstrumentRelation.knowledge_base_id.in_(kb_ids),
        LegalInstrumentRelation.review_status == "verified",
        LegalInstrumentRelation.relation == "AMENDS",
        Document.deleted_at.is_(None), Document.status == "completed",
    )
    if family_id:
        target_instrument = aliased(LegalInstrument)
        query = query.join(target_instrument, target_instrument.id == LegalInstrumentRelation.target_instrument_id).filter(
            target_instrument.family_id == family_id
        )
    return query.order_by(LegalInstrument.version_date.desc(), Document.created_at.desc()).all()


def _latest_amendment_for_provision(rows: list[tuple[LegalInstrumentRelation, LegalInstrument, Document]], number: str) -> tuple[LegalInstrument, Document] | None:
    """Pick the most recent verified act, from an already-fetched and
    family-scoped relation list, that amended a given section number.

    Lets an answer served from the consolidated latest version state which
    amending act produced a provision's current wording (e.g. มาตรา 104's text
    comes from พ.ร.บ.แก้ไขฯ ฉบับที่ 15), which the instrument-level status alone
    does not convey.
    """
    for relation, instrument, document in rows:
        if provision_number_matches(relation.target_provision, number):
            return instrument, document
    return None


def _persist_legal_clause_result(db: Session, *, query: str, kb_ids: list[str], token_id: str | None, intent: str,
                                 detail: str, rows: list[tuple[DocumentChunk, Document, LegalInstrument]],
                                 attribute_amendments: bool = False) -> dict:
    # Defense-in-depth: drop verbatim-duplicate clauses (a consolidated file
    # repeats identical commencement text) and cap the result, without collapsing
    # a long provision whose distinct chunks legitimately share a section number.
    deduped: list[tuple[DocumentChunk, Document, LegalInstrument]] = []
    seen_excerpts: set[str] = set()
    for chunk, document, instrument in rows:
        key = (chunk.content or "").strip()
        if key in seen_excerpts:
            continue
        seen_excerpts.add(key)
        deduped.append((chunk, document, instrument))
    rows = deduped[:6]
    sources = [{
        "citation_id": f"S{index}", "document_id": document.id, "chunk_id": chunk.id,
        "title": document.title or document.original_filename, "section_label": chunk.section_label,
        "excerpt": chunk.content, "relevance": 1.0 / index, "document_status": instrument.status,
        "document_class": instrument.document_class, "version_role": instrument.version_role,
        "source_uri": instrument.source_uri, "source_reference": instrument.source_reference,
        "legal_label": instrument.official_title or document.title or document.original_filename,
    } for index, (chunk, document, instrument) in enumerate(rows, 1)]
    # The raw clause is a factual legal claim. Keep its citation adjacent to
    # the text so callers do not need to infer which trailing source applies.
    answer = "\n\n".join(f"{source['excerpt']}\n[{source['citation_id']}]" for source in sources)
    attribution_lines: list[str] = []
    if attribute_amendments and sources:
        family_id = rows[0][2].family_id if rows else None
        amendment_rows = _verified_amendment_relations(db, kb_ids, family_id)
        for number in dict.fromkeys(chunk.section_number for chunk, _, _ in rows if chunk.section_number):
            found = _latest_amendment_for_provision(amendment_rows, number)
            if not found:
                continue
            amending_instrument, amending_document = found
            citation_id = f"S{len(sources) + 1}"
            amend_title = amending_instrument.official_title or amending_document.title or amending_document.original_filename
            sources.append({
                "citation_id": citation_id, "document_id": amending_document.id,
                "title": amending_document.title or amending_document.original_filename,
                "section_label": f"มาตรา {number}", "excerpt": f"มาตรา {number} แก้ไขเพิ่มเติมโดย {amend_title}",
                "relevance": 0.6, "document_status": amending_instrument.status,
                "document_class": amending_instrument.document_class, "version_role": amending_instrument.version_role,
                "version_date": amending_instrument.version_date.isoformat() if amending_instrument.version_date else None,
                "source_uri": amending_instrument.source_uri, "source_reference": amending_instrument.source_reference,
                "legal_label": amend_title,
            })
            attribution_lines.append(f"ข้อความปัจจุบันของมาตรา {number} แก้ไขเพิ่มเติมโดย {amend_title} [{citation_id}]")
    if sources:
        answer = "คำตอบนี้อ้างอิงฉบับปรับปรุงล่าสุดที่มีผลบังคับใช้ในคลังข้อมูล\n" + answer
        if attribution_lines:
            answer += "\n\n" + "\n".join(attribution_lines)
        answer += "\n\nรายละเอียดแหล่งอ้างอิง:\n" + _render_citation_details(sources)
    else:
        answer = "ไม่พบข้อบทที่ตรงกับเงื่อนไขในคลังข้อมูลที่เลือก"
    result = {"status": "success", "result_id": "", "answer": answer, "insufficient_evidence": not bool(sources), "sources": sources,
              "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": "legal_clause_registry",
              "retrieval_plan": {"version": 1, "intent": intent, "planner_source": "rules", "channels": ["legal_registry", "document_chunks"]},
              "retrieval_trace": [{"channel": "document_chunks", "system": "PostgreSQL legal sections", "status": "used", "result_count": len(sources), "duration_ms": 0, "detail": detail}],
              **_trace_preview_fields(query, answer, sources)}}
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result


def _sole_latest_legal_instrument(db: Session, kb_ids: list[str]) -> tuple[Document, LegalInstrument] | None:
    rows = db.query(Document, LegalInstrument).join(LegalInstrument, LegalInstrument.document_id == Document.id).filter(
        Document.knowledge_base_id.in_(kb_ids), Document.deleted_at.is_(None), Document.status == "completed",
        LegalInstrument.version_role == "latest_consolidated",
    ).all()
    return rows[0] if len(rows) == 1 else None


# A consolidated Thai code file concatenates several sub-works and repeats
# section numbers across them: the enabling act (พระราชบัญญัติให้ใช้…), the code
# proper, and the trailing commencement/หมายเหตุ clauses of every amendment act.
# ``ในประมวลกฎหมายนี้`` opens the code's own definitional section 1; the amendment
# footers each carry the explanatory ``หมายเหตุ :- เหตุผลในการประกาศ`` line.
_CODE_BODY_START_MARKER = "ในประมวลกฎหมายนี้"
_AMENDMENT_NOTE_MARKER = "หมายเหตุ :- เหตุผลในการประกาศ"
# Reusable legal concepts that occur verbatim in questions; used to rank clauses
# lexically when a bare section number collides across sub-works.
_LEGAL_CONCEPT_TERMS = ("ค่าธรรมเนียม", "จดทะเบียน", "นิติกรรม", "โอน", "กรรมสิทธิ์", "สิทธิครอบครอง",
                        "ราคาประเมิน", "ทุนทรัพย์", "ใบอนุญาต", "ค่าตอบแทน", "ที่ดินของรัฐ", "หวงห้าม",
                        "องค์การบริหารส่วนจังหวัด", "องค์กรปกครองส่วนท้องถิ่น", "รายได้")


def segment_document_subworks(chunks: list[DocumentChunk]) -> dict[str, str]:
    """Classify each chunk of a consolidated legal document into a sub-work:
    ``enabling_act`` (the พระราชบัญญัติให้ใช้… preamble act), ``code_body`` (the
    code proper), or ``amendment_note`` (the commencement/หมายเหตุ clauses that a
    consolidation appends verbatim for every amending act).

    One consolidated file repeats section numbers across these parts (two
    ``มาตรา 9``, dozens of ``มาตรา 2``), so builders that filter by section number
    alone need this to keep an enabling-act or amendment-note section out of an
    answer about the code proper. Detection is deterministic, ordered by
    ``chunk_index``, and fails safe: a document with no recognizable code-body
    boundary returns every chunk as ``unknown`` so callers keep prior behavior.
    """
    ordered = sorted(chunks, key=lambda chunk: chunk.chunk_index)
    stage = "enabling_act"
    code_started = False
    result: dict[str, str] = {}
    for chunk in ordered:
        text = chunk.content or ""
        # Amendment notes only ever follow the code body itself, never the
        # enabling act; requiring code_body to have already started rules out
        # an incidental match of the marker phrase while still in the
        # enabling act (e.g. a cross-reference in its preamble text) from
        # short-circuiting classification for the rest of the document.
        if stage == "code_body" and _AMENDMENT_NOTE_MARKER in text:
            stage = "amendment_note"
        # ในประมวลกฎหมายนี้ is definitionally the opening of a Thai code's
        # มาตรา 1; requiring the section number match too rules out an
        # incidental mention of the phrase elsewhere in the enabling act.
        elif not code_started and _CODE_BODY_START_MARKER in text and chunk.section_number == "1":
            stage = "code_body"
            code_started = True
        result[chunk.id] = stage
    if not code_started:
        return {chunk.id: "unknown" for chunk in ordered}
    return result


def _rank_chunks_by_query(query: str, chunks: list[DocumentChunk]) -> list[tuple[DocumentChunk, int]]:
    """Order chunks by lexical overlap with the question, stable on chunk order.

    Thai has no word spaces, so we match the salient phrases the user did
    separate plus a reusable legal-concept list; a section number that collides
    across sub-works then resolves to the clause the question is really about.
    Returns (chunk, score) pairs, highest first, so a caller that must fail
    closed on no match can filter out zero-score chunks itself.
    """
    value = (query or "").casefold()
    phrases = {phrase for phrase in re.split(r"[\s?,.()ฯ]+", value) if len(phrase) >= 3}
    terms = {term for term in (phrases | set(_LEGAL_CONCEPT_TERMS)) if term and term in value}

    def score(chunk: DocumentChunk) -> int:
        text = (chunk.content or "").casefold()
        return sum(1 for term in terms if term in text)

    scored = [(chunk, score(chunk)) for chunk in chunks]
    scored.sort(key=lambda item: (-item[1], item[0].chunk_index))
    return scored


def is_legal_commencement_lookup(query: str) -> bool:
    value = (query or "").casefold()
    return "เริ่มใช้บังคับ" in value and not _requested_amendment_number(query)


def build_legal_commencement_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    current = _sole_latest_legal_instrument(db, kb_ids)
    rows: list[tuple[DocumentChunk, Document, LegalInstrument]] = []
    if current:
        document, instrument = current
        all_chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).order_by(DocumentChunk.chunk_index).all()
        sub_works = segment_document_subworks(all_chunks)
        # Section 2/3 collide across dozens of amendment footers; keep only the
        # real commencement sections and prefer the clause stating when the code
        # itself came into force over the enabling act's own commencement.
        candidates = [chunk for chunk in all_chunks
                      if chunk.section_number in ("2", "3") and sub_works.get(chunk.id) != "amendment_note"]
        focused = [chunk for chunk in candidates
                   if "ประมวลกฎหมายที่ดิน" in (chunk.content or "") and "ให้ใช้บังคับ" in (chunk.content or "")]
        rows = [(chunk, document, instrument) for chunk in (focused or candidates)[:2]]
    return _persist_legal_clause_result(db, query=query, kb_ids=kb_ids, token_id=token_id, intent="legal_commencement_rule",
                                        detail="current legal commencement clauses", rows=rows)


def is_legal_document_copy_lookup(query: str) -> bool:
    value = (query or "").casefold()
    return ("โฉนดที่ดิน" in value or "หนังสือรับรองการทำประโยชน์" in value) and any(
        term in value for term in ("กี่ฉบับ", "ใครเก็บ", "ต้นฉบับ", "จัดเก็บ")
    )


def build_legal_document_copy_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    current = _sole_latest_legal_instrument(db, kb_ids)
    rows: list[tuple[DocumentChunk, Document, LegalInstrument]] = []
    if current:
        document, instrument = current
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document.id,
            DocumentChunk.content.ilike("%คู่ฉบับ%"),
            DocumentChunk.content.ilike("%โฉนดที่ดิน%"),
            DocumentChunk.content.ilike("%หนังสือรับรองการทำประโยชน์%"),
        ).order_by(DocumentChunk.chunk_index).all()
        rows = [(chunk, document, instrument) for chunk in chunks[:2]]
    return _persist_legal_clause_result(db, query=query, kb_ids=kb_ids, token_id=token_id, intent="legal_document_copy_rule",
                                        detail="complete document-copy clause", rows=rows)


def build_default_current_provision_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    """Serve a cited complete section from the sole current consolidation.

    Semantic retrieval remains the normal path for multi-instrument scopes;
    this deterministic path is safe only where the registry identifies exactly
    one current consolidated expression.
    """
    current = _sole_latest_legal_instrument(db, kb_ids)
    refs = parse_provision_refs(query)
    rows: list[tuple[DocumentChunk, Document, LegalInstrument]] = []
    if current and refs:
        document, instrument = current
        chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).order_by(DocumentChunk.chunk_index).all()
        sub_works = segment_document_subworks(chunks)
        matching = [chunk for chunk in chunks if any(provision_number_matches(ref["number"], chunk.section_number) for ref in refs)]
        # A bare number matches มาตรา and หมวด alike; when the question names
        # มาตรา, keep those and drop a same-numbered chapter heading.
        requested_kinds = {ref["kind"] for ref in refs}
        matching = [chunk for chunk in matching if chunk.section_kind in requested_kinds] or matching
        # The same section number appears in the enabling act and in amendment
        # footers; keep the code proper so e.g. "มาตรา 9" resolves to the code's
        # provision, not the enabling act's. Rank collisions by the question.
        matching = [chunk for chunk in matching if sub_works.get(chunk.id) != "amendment_note"] or matching
        code_body = [chunk for chunk in matching if sub_works.get(chunk.id) == "code_body"]
        pool = code_body or matching
        # Pick the provisions the question is really about, then include every
        # chunk of those sections in reading order so a provision split across
        # chunks (e.g. มาตรา 9/1) is answered in full, not truncated.
        top_numbers = {chunk.section_number for chunk, _ in _rank_chunks_by_query(query, pool)[:2]}
        selected = [chunk for chunk in sorted(pool, key=lambda chunk: chunk.chunk_index)
                    if chunk.section_number in top_numbers]
        rows = [(chunk, document, instrument) for chunk in selected[:4]]
    return _persist_legal_clause_result(db, query=query, kb_ids=kb_ids, token_id=token_id, intent="legal_current_provision",
                                        detail="default-current complete provision", rows=rows, attribute_amendments=True)


def is_default_current_legal_lookup(query: str) -> bool:
    value = (query or "").casefold()
    return any(term in value for term in ("ปัจจุบัน", "ล่าสุด")) and any(
        term in value for term in ("กฎหมาย", "หลักเกณฑ์", "ค่าธรรมเนียม", "สิทธิ", "หน้าที่")
    )


def build_default_current_legal_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    """Lexically locate the best clause in a sole current legal expression.

    Used only for default-current questions that do not name a provision.  It
    prevents an old amendment from winning vector similarity while remaining
    reusable for any legal work with a designated latest consolidation.
    """
    current = _sole_latest_legal_instrument(db, kb_ids)
    rows: list[tuple[DocumentChunk, Document, LegalInstrument]] = []
    if current:
        document, instrument = current
        chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).all()
        ranked = _rank_chunks_by_query(query, chunks)
        # Fail closed to an empty answer when nothing in the document actually
        # matches the question, rather than citing two arbitrary chunks.
        rows = [(chunk, document, instrument) for chunk, score in ranked[:2] if score > 0]
    return _persist_legal_clause_result(db, query=query, kb_ids=kb_ids, token_id=token_id, intent="legal_default_current",
                                        detail="sole current legal expression lexical clause", rows=rows, attribute_amendments=True)


_PROVENANCE_ACTION_LABELS = {"REPEALS": "ยกเลิก", "SUPERSEDES": "แทนที่", "AMENDS": "แก้ไขเพิ่มเติม"}
_TERMINAL_PROVENANCE_RELATIONS = {"REPEALS", "SUPERSEDES"}


def build_legal_provenance_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    """Return verified amendment provenance without relying on semantic ranking.

    Relations are created only from explicit change clauses during legal extraction,
    so this path deliberately fails closed if the registry has no matching edge.
    """
    refs = parse_provision_refs(query)
    relation_rows = db.query(LegalInstrumentRelation, LegalInstrument, Document).join(
        LegalInstrument, LegalInstrument.id == LegalInstrumentRelation.source_instrument_id
    ).join(Document, Document.id == LegalInstrument.document_id).filter(
        LegalInstrumentRelation.knowledge_base_id.in_(kb_ids),
        LegalInstrumentRelation.review_status == "verified",
        LegalInstrumentRelation.relation.in_(("AMENDS", "REPEALS", "SUPERSEDES")),
        Document.deleted_at.is_(None), Document.status == "completed",
    ).order_by(LegalInstrument.version_date.desc(), Document.created_at.desc()).all()
    target_ids = {row[0].target_instrument_id for row in relation_rows if row[0].target_instrument_id}
    target_rows = db.query(LegalInstrument, Document).join(Document, Document.id == LegalInstrument.document_id).filter(
        LegalInstrument.id.in_(target_ids), Document.deleted_at.is_(None), Document.status == "completed",
    ).all() if target_ids else []
    targets = {instrument.id: (instrument, document) for instrument, document in target_rows}
    sources, statements = [], []
    chunks_by_document: dict[str, list[DocumentChunk]] = {}

    def emit_relation(relation: LegalInstrumentRelation, instrument: LegalInstrument, document: Document, provision: str) -> None:
        citation_id = f"S{len(sources) + 1}"
        title = document.title or document.original_filename
        action = _PROVENANCE_ACTION_LABELS.get(relation.relation, "แก้ไขเพิ่มเติม")
        statement_citations = [citation_id]
        sources.append({
            "citation_id": citation_id, "document_id": document.id, "title": title,
            "section_label": f"มาตรา {provision}", "excerpt": relation.evidence_quote or f"{action} {provision}",
            "relevance": 1.0, "document_status": instrument.status, "document_class": instrument.document_class,
            "version_role": instrument.version_role, "version_date": instrument.version_date.isoformat() if instrument.version_date else None,
            "source_uri": instrument.source_uri, "source_reference": instrument.source_reference,
            "legal_label": instrument.official_title or title,
        })
        target = targets.get(relation.target_instrument_id)
        if target:
            target_instrument, target_document = target
            # Several relations (a compound question naming multiple
            # provisions, or several repeal/amend rows on the same act) can
            # point at the same target document; fetch its chunks once.
            if target_document.id not in chunks_by_document:
                chunks_by_document[target_document.id] = db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == target_document.id,
                ).order_by(DocumentChunk.chunk_index).all()
            target_chunks = chunks_by_document[target_document.id]
            target_chunk = next((chunk for chunk in target_chunks if provision_number_matches(provision, chunk.section_number)), None)
            target_citation_id = f"S{len(sources) + 1}"
            sources.append({
                "citation_id": target_citation_id, "document_id": target_document.id,
                "chunk_id": target_chunk.id if target_chunk else None,
                "title": target_document.title or target_document.original_filename,
                "section_label": target_chunk.section_label if target_chunk else f"มาตรา {provision}",
                "excerpt": (target_chunk.content if target_chunk else f"ข้อบทปลายทาง {provision}"),
                "relevance": 0.95, "document_status": target_instrument.status,
                "document_class": target_instrument.document_class, "version_role": target_instrument.version_role,
                "version_date": target_instrument.version_date.isoformat() if target_instrument.version_date else None,
                "source_uri": target_instrument.source_uri, "source_reference": target_instrument.source_reference,
                "legal_label": target_instrument.official_title or target_document.title or target_document.original_filename,
            })
            statement_citations.append(target_citation_id)
        statements.append(f"มาตรา {provision} {action}โดย {instrument.official_title or title} {' '.join(f'[{item}]' for item in statement_citations)}")

    # Answer each provision the user asked about, so a compound question like
    # "มาตรา 96 ทวิ และมาตรา 96 ตรี" reports on both instead of only the first.
    for ref in refs:
        ref_rows = [row for row in relation_rows if provision_number_matches(row[0].target_provision, ref["number"])]
        deduped: list[tuple[LegalInstrumentRelation, LegalInstrument, Document]] = []
        seen: set[tuple[str, str]] = set()
        for relation, instrument, document in ref_rows:
            dedupe_key = (instrument.document_id, relation.target_provision or ref["number"])
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            deduped.append((relation, instrument, document))
        if not deduped:
            statements.append(f"ไม่พบความสัมพันธ์การแก้ไขหรือยกเลิกที่ตรวจสอบแล้วสำหรับมาตรา {ref['number']} ในคลังข้อมูลที่เลือก")
            continue
        # A repeal or supersession is terminal for a provision, matching how
        # validate_legal_evidence treats REPEALS and SUPERSEDES identically as
        # overriding relations (see legal_registry.py's resolve_instrument_statuses,
        # which names that function the single owner of this precedence rule).
        # Once มาตรา N of the code is repealed/superseded, an "amend" of มาตรา N
        # belongs to a same-numbered provision in another work (e.g. the
        # พ.ร.บ.ให้ใช้ฯ), not the terminated section. Lead with the terminal
        # relation and flag any amend separately instead of emitting a
        # contradictory line.
        terminal = [row for row in deduped if row[0].relation in _TERMINAL_PROVENANCE_RELATIONS]
        amends = [row for row in deduped if row[0].relation == "AMENDS"]
        if terminal:
            for relation, instrument, document in terminal:
                emit_relation(relation, instrument, document, relation.target_provision or ref["number"])
            if amends:
                other_acts = ", ".join(dict.fromkeys(
                    instrument.official_title or (document.title or document.original_filename)
                    for _, instrument, document in amends))
                statements.append(
                    f"หมายเหตุ: พบความสัมพันธ์ 'แก้ไขเพิ่มเติม' มาตรา {ref['number']} โดย {other_acts} "
                    f"ซึ่งอ้างถึงมาตรา {ref['number']} ในกฎหมายคนละฉบับ ไม่ใช่บทที่ถูกยกเลิกหรือถูกแทนที่ข้างต้น")
        else:
            for relation, instrument, document in amends:
                emit_relation(relation, instrument, document, relation.target_provision or ref["number"])
    if statements:
        answer = "\n".join(statements) + "\n\nรายละเอียดแหล่งอ้างอิง:\n" + _render_citation_details(sources)
    else:
        answer = "ไม่พบความสัมพันธ์การแก้ไขหรือยกเลิกที่ตรวจสอบแล้วสำหรับมาตราที่ระบุในคลังข้อมูลที่เลือก"
    result = {
        "status": "success", "result_id": "", "answer": answer, "insufficient_evidence": not bool(sources),
        "sources": sources, "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": "legal_provenance_registry",
        "retrieval_plan": {"version": 1, "intent": "legal_provenance_lookup", "planner_source": "rules", "channels": ["legal_registry"]},
        "retrieval_trace": [{"channel": "legal_registry", "system": "PostgreSQL legal registry", "status": "used", "result_count": len(sources), "duration_ms": 0, "detail": "verified amendment provenance lookup"}],
        **_trace_preview_fields(query, answer, sources)},
    }
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result


def legal_scope_gap_response(query: str) -> str:
    """Make legal no-evidence responses actionable while preserving fail-closed behavior."""
    value = (query or "").casefold()
    if "คำพิพากษา" in value or "ฎีกา" in value:
        required = "คำพิพากษาศาลหรือฐานข้อมูลคำพิพากษาที่มีเลขคดีและข้อความเต็ม"
    elif "ประกาศ" in value or "ระเบียบ" in value:
        required = "ประกาศหรือระเบียบของหน่วยงานที่เกี่ยวข้อง"
    else:
        required = "เอกสารกฎหมายหรือเอกสารทางการที่ครอบคลุมประเด็นนี้"
    return f"ไม่พบหลักฐานยืนยันคำตอบในขอบเขตของคลังความรู้ที่เลือก จึงไม่ควรสรุปข้อเท็จจริงเพิ่มเติม ต้องเพิ่ม {required} ก่อนจึงจะตอบได้อย่างตรวจสอบย้อนกลับได้"


def _persist_scope_gap_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None) -> dict:
    answer = legal_scope_gap_response(query)
    result = {
        "status": "success", "result_id": "", "answer": answer, "insufficient_evidence": True,
        "sources": [], "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": "legal_scope_control",
        "retrieval_plan": {"version": 1, "intent": "legal_out_of_scope", "planner_source": "rules", "channels": ["legal_registry"]},
        "retrieval_trace": [{"channel": "legal_registry", "system": "PostgreSQL legal registry", "status": "used", "result_count": 0, "duration_ms": 0, "detail": "out-of-scope document-type check"}],
        **_trace_preview_fields(query, answer, [])},
    }
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result


def has_court_decision_evidence(db: Session, kb_ids: list[str]) -> bool:
    """Detect court-decision material from document type, template, and indexed metadata."""
    if not kb_ids:
        return False
    marker = "%คำพิพากษา%"
    row = db.query(Document.id).filter(
        Document.knowledge_base_id.in_(kb_ids), Document.deleted_at.is_(None), Document.status == "completed",
        or_(
            Document.document_type.ilike("%judgment%"),
            Document.metadata_template_name.ilike(marker),
            Document.metadata_search_text.ilike(marker),
            cast(Document.document_metadata, Text).ilike(marker),
            Document.title.ilike(marker),
        ),
    ).first()
    return row is not None


def has_active_query_filters(query_filters=None) -> bool:
    return bool(query_filters and (
        getattr(query_filters, "published_from", None)
        or getattr(query_filters, "published_to", None)
        or getattr(query_filters, "as_of_date", None)
        or getattr(query_filters, "include_historical", False)
        or getattr(query_filters, "metadata", None)
    ))


def allows_direct_registry_path(query_filters=None) -> bool:
    """Registry shortcuts cannot silently bypass caller-supplied retrieval filters."""
    return not has_active_query_filters(query_filters)


def allows_default_current_direct_path(query: str, query_filters=None) -> bool:
    """Current-only shortcuts must never override an explicit historical request."""
    if not allows_direct_registry_path(query_filters):
        return False
    value = (query or "").casefold()
    return not bool(_requested_amendment_number(query) or re.search(r"พ\.ศ\.?\s*[0-9๐-๙]{4}", value))


def summarize_document_inventory(db: Session, kb_ids: list[str], *, scope: str = "all",
                                 include_documents: bool = True, max_documents: int = 500) -> dict:
    """Return an authoritative, non-LLM document and legal-registry summary.

    ``scope=all`` means every non-deleted document in the token's active KB
    scope. ``scope=current`` excludes legal instruments marked superseded,
    repealed, or not-yet-effective; it never includes deleted documents.
    """
    current_statuses = {"in_force", "amended", "unknown"}
    base_query = db.query(Document, LegalInstrument).outerjoin(
        LegalInstrument, LegalInstrument.document_id == Document.id
    ).filter(Document.knowledge_base_id.in_(kb_ids), Document.deleted_at.is_(None))
    if scope == "current":
        base_query = base_query.filter(or_(Document.document_type.notin_(LEGAL_DOCUMENT_TYPES), LegalInstrument.status.in_(current_statuses)))

    total_documents = base_query.with_entities(func.count(Document.id)).scalar() or 0
    type_rows = base_query.with_entities(Document.document_type, func.count(Document.id)).group_by(Document.document_type).all()
    kind_rows = base_query.filter(LegalInstrument.id.is_not(None)).with_entities(LegalInstrument.kind, func.count(Document.id)).group_by(LegalInstrument.kind).all()
    class_rows = base_query.filter(LegalInstrument.document_class.is_not(None)).with_entities(LegalInstrument.document_class, func.count(Document.id)).group_by(LegalInstrument.document_class).all()
    role_rows = base_query.filter(LegalInstrument.version_role.is_not(None)).with_entities(LegalInstrument.version_role, func.count(Document.id)).group_by(LegalInstrument.version_role).all()
    rows = base_query.order_by(Document.title, Document.original_filename, Document.id).limit(max_documents if include_documents else 0).all() if include_documents else []

    type_counts = {key: count for key, count in type_rows}
    kind_counts = {key: count for key, count in kind_rows}
    class_counts = {key: count for key, count in class_rows}
    role_counts = {key: count for key, count in role_rows}
    documents: list[dict] = []
    for document, instrument in rows:
        if include_documents and len(documents) < max_documents:
            documents.append({
                "citation_id": f"D{len(documents) + 1}",
                "document_id": document.id,
                "title": document.title or document.original_filename,
                "original_filename": document.original_filename,
                "document_type": document.document_type,
                "document_type_label": DOCUMENT_TYPE_LABELS.get(document.document_type, document.document_type),
                "status": document.status,
                "published_at": document.published_at.isoformat() if document.published_at else None,
                "legal_kind": instrument.kind if instrument else None,
                "legal_kind_label": LEGAL_KIND_LABELS.get(instrument.kind, instrument.kind) if instrument else None,
                "document_class": instrument.document_class if instrument else None,
                "document_class_label": DOCUMENT_CLASS_LABELS.get(instrument.document_class, instrument.document_class) if instrument and instrument.document_class else None,
                "version_role": instrument.version_role if instrument else None,
                "version_role_label": VERSION_ROLE_LABELS.get(instrument.version_role, instrument.version_role) if instrument and instrument.version_role else None,
                "legal_status": instrument.status if instrument else None,
                "official_number": instrument.official_number if instrument else None,
            })

    groups: list[dict] = []
    for key, count in sorted(type_counts.items()):
        groups.append({"dimension": "document_type", "key": key, "label": DOCUMENT_TYPE_LABELS.get(key, key), "count": count})
    for key, count in sorted(kind_counts.items()):
        groups.append({"dimension": "legal_kind", "key": key, "label": LEGAL_KIND_LABELS.get(key, key), "count": count})
    for key, count in sorted(class_counts.items()):
        groups.append({"dimension": "document_class", "key": key, "label": DOCUMENT_CLASS_LABELS.get(key, key), "count": count})
    for key, count in sorted(role_counts.items()):
        groups.append({"dimension": "version_role", "key": key, "label": VERSION_ROLE_LABELS.get(key, key), "count": count})
    return {
        "total_documents": total_documents,
        "total_legal_instruments": sum(count for _, count in kind_rows),
        "scope": scope,
        "groups": groups,
        "documents": documents,
        "documents_truncated": include_documents and total_documents > len(documents),
        "source_of_truth": "PostgreSQL document and legal registry",
    }


def _inventory_sources(inventory: dict) -> list[dict]:
    """Create compact, citation-friendly evidence for an inventory summary."""
    grouped = [group for group in inventory["groups"] if group["dimension"] in {"version_role", "document_type"}]
    sources = []
    for index, group in enumerate(grouped, 1):
        sources.append({
            "citation_id": f"I{index}",
            "document_id": None,
            "title": group["label"],
            "excerpt": f"{group['label']}: {group['count']} รายการ ({group['dimension']})",
            "relevance": 1.0 / index,
            "inventory_dimension": group["dimension"],
            "inventory_key": group["key"],
            "inventory_count": group["count"],
            "document_ids": [item["document_id"] for item in inventory["documents"]
                             if (group["dimension"] == "document_type" and item["document_type"] == group["key"])
                             or (group["dimension"] == "version_role" and item["version_role"] == group["key"])],
        })
    if not sources and inventory["total_documents"]:
        sources.append({"citation_id": "I1", "document_id": None, "title": "Document inventory",
                        "excerpt": f"พบเอกสารทั้งหมด {inventory['total_documents']} รายการ",
                        "relevance": 1.0, "inventory_count": inventory["total_documents"]})
    return sources


def _inventory_answer(inventory: dict, sources: list[dict]) -> str:
    if not inventory["total_documents"]:
        return "ไม่พบเอกสารใน Knowledge Base ที่อยู่ในขอบเขตของ MCP key [I1]" if sources else "ไม่พบเอกสารใน Knowledge Base ที่อยู่ในขอบเขตของ MCP key"
    lines = [f"พบเอกสารทั้งหมด {inventory['total_documents']} รายการ (ขอบเขต: {inventory['scope']})"]
    for dimension in ("version_role",):
        groups = [group for group in inventory["groups"] if group["dimension"] == dimension]
        if groups:
            label = "ประเภท/ชั้นของเอกสารกฎหมาย"
            lines.append(f"{label}: " + ", ".join(f"{group['label']} {group['count']} รายการ [I{sources.index(next(source for source in sources if source.get('inventory_key') == group['key'] and source.get('inventory_dimension') == dimension)) + 1}]" for group in groups))
    return "\n".join(lines)


def build_document_inventory_result(db: Session, query: str, kb_ids: list[str], token_id: str | None = None,
                                    query_filters=None, *, scope: str | None = None,
                                    include_documents: bool = True, max_documents: int = 500) -> dict:
    """Build a saved, deterministic result for REST/MCP inventory queries."""
    resolved_scope = scope or document_inventory_scope(query, query_filters)
    inventory = summarize_document_inventory(db, kb_ids, scope=resolved_scope,
                                             include_documents=include_documents, max_documents=max_documents)
    sources = _inventory_sources(inventory)
    answer = _inventory_answer(inventory, sources)
    plan = {"version": 1, "intent": "document_inventory", "planner_source": "rules",
            "rationale": "deterministic document/legal registry aggregation", "channels": ["document_registry"],
            "max_sources": len(sources), "graph_depth": 0, "graph_scope": "none", "entity_subjects": [],
            "document_identifiers": [], "published_from": None, "published_to": None, "as_of_date": None,
            "include_historical": resolved_scope == "all", "rerank_enabled": False}
    trace = [{"channel": "document_registry", "system": "PostgreSQL document and legal registry",
              "status": "used", "result_count": inventory["total_documents"], "duration_ms": 0,
              "detail": f"scope={resolved_scope}; deterministic count and grouping"}]
    result = {"status": "success", "result_id": "", "answer": answer,
              "insufficient_evidence": not bool(inventory["total_documents"]), "sources": sources,
              "documents": inventory["documents"], "total_documents": inventory["total_documents"],
              "total_legal_instruments": inventory["total_legal_instruments"], "groups": inventory["groups"],
              "scope": resolved_scope, "documents_truncated": inventory["documents_truncated"],
              "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": "document_inventory",
                           "retrieval_plan": plan, "retrieval_trace": trace, "source_of_truth": inventory["source_of_truth"],
                           "query_preview": query[:500] if get_settings().log_query_text else None,
                           "query_length": len(query), "query_sha256": hashlib.sha256(query.encode("utf-8")).hexdigest(),
                           "answer_preview": answer[:800] if get_settings().log_query_text else None,
                           "citation_ids": [source["citation_id"] for source in sources],
                           "filter_summary": query_filters.model_dump(mode="json") if query_filters else {},
                           "response_summary": {"status": "success", "insufficient_evidence": not bool(inventory["total_documents"]),
                                                "source_count": len(sources), "entity_count": 0, "relationship_count": 0,
                                                "total_documents": inventory["total_documents"]}}}
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result


def create_document_job(db: Session, knowledge_base_id: str, upload, title: str | None = None, document_type: str = "general", published_at=None,
                        metadata_template: dict | None = None, document_metadata: dict | None = None) -> tuple[Document, ProcessingJob]:
    if document_type not in DOCUMENT_TYPES:
        raise ValueError("DOCUMENT_TYPE_INVALID")
    path, stored, size, checksum, mime = store_upload(upload, knowledge_base_id)
    duplicate = db.query(Document).filter_by(knowledge_base_id=knowledge_base_id, checksum_sha256=checksum).filter(Document.deleted_at.is_(None)).first()
    if duplicate:
        Path(path).unlink(missing_ok=True)
        raise ValueError("FILE_DUPLICATE")
    fields = normalize_field_definitions((metadata_template or {}).get("fields") or [])
    values = document_metadata or {}
    doc = Document(knowledge_base_id=knowledge_base_id, original_filename=upload.filename or stored,
                   stored_filename=stored, storage_path=path, file_size=size, checksum_sha256=checksum,
                   mime_type=mime, title=title or Path(upload.filename or stored).stem,
                   document_type=document_type, published_at=published_at, status="queued",
                   metadata_template_id=(metadata_template or {}).get("id"),
                   metadata_template_name=(metadata_template or {}).get("name"),
                   metadata_template_version=(metadata_template or {}).get("version"),
                   metadata_template_fields=fields,
                   document_metadata=values,
                   metadata_search_text=metadata_search_text(fields, values))
    db.add(doc)
    try:
        db.flush()
    except IntegrityError as exc:
        # F8 race guard: two concurrent uploads of the same file — the
        # partial unique index is the real guard; surface the same 409 the
        # non-concurrent path returns (the text path already did this).
        # Only a unique-checksum violation maps to FILE_DUPLICATE; any other
        # integrity failure (FK, NOT NULL) must surface as itself (review M1).
        # str(exc) carries the constraint/column name on both SQLite and
        # Postgres (constraint_name lives in exc.orig.diag on PG only).
        detail = str(exc)
        if "uq_document_checksum" not in detail and "checksum_sha256" not in detail:
            raise
        db.rollback()
        Path(path).unlink(missing_ok=True)
        raise ValueError("FILE_DUPLICATE") from None
    job = ProcessingJob(document_id=doc.id, knowledge_base_id=knowledge_base_id)
    db.add(job); db.commit(); db.refresh(doc); db.refresh(job)
    return doc, job


def _legacy_extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md", ".csv", ".json"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if ext in {".html", ".htm"}:
        return BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser").get_text(" ", strip=True)
    if ext == ".docx":
        return "\n".join(p.text for p in WordDocument(path).paragraphs)
    if ext == ".pdf":
        # PDF extraction is owned by the anydoc pipeline (Rust + OCR chain).
        # Without the wheel there is no safe local PDF fallback left: pypdf
        # was removed with the legacy pipeline it served.
        raise RuntimeError("ANYDOC_UNAVAILABLE")
    raise RuntimeError("FILE_TYPE_NOT_SUPPORTED")


def _markitdown_extract(path: Path) -> str:
    """Convert a validated local upload to Markdown without plugins, OCR, LLMs, or cloud services."""
    result = MarkItDown(enable_builtins=True, enable_plugins=False).convert_local(path)
    return result.text_content.strip()


def _has_meaningful_text(text: str) -> bool:
    return len(re.sub(r"[\W_]+", "", text, flags=re.UNICODE)) >= 20


def _thai_ratio(text: str) -> float:
    """Share of characters that are Thai. Mojibake of Thai text renders as
    Latin/graphic glyphs, so a Thai document whose ratio is ~0 carries no
    usable text layer even though it passes a bare character count."""
    if not text:
        return 0.0
    thai = sum(1 for ch in text if "\u0e00" <= ch <= "\u0e7f")
    return thai / len(text)


def _non_ascii_ratio(text: str) -> float:
    """Share of characters outside printable ASCII. Genuine English (or any
    plain-ASCII) document stays near zero; Thai text read through a Latin
    glyph table turns almost every character into an extended one."""
    if not text:
        return 0.0
    plain = sum(1 for ch in text if 0x20 <= ord(ch) < 0x7f or ch in "\n\t")
    return 1.0 - plain / len(text)


# A Thai PDF whose text layer carries no Thai at all but plenty of extended
# Latin glyphs is presumed garbled mojibake. Plain-ASCII documents (English)
# are unaffected: they carry almost no extended characters. A short extract
# (headers, page numbers) is too small to judge, so only longer layers are
# checked.
_MIN_TEXT_FOR_LANGUAGE_CHECK = 200
_MAX_THAI_RATIO_FOR_GARBLED = 0.02
_MIN_NON_ASCII_RATIO_FOR_GARBLED = 0.30

# Broken PDF font mappings observed in the wild: TIS-620 bytes surfaced
# through a mac-roman glyph table. Round-tripping the visible text through
# that pair recovers the original deterministically, without OCR.
_CODEC_RECOVERY_PAIRS: tuple[tuple[str, str], ...] = (
    ("mac_roman", "tis-620"),
    ("cp1252", "utf-8"),
)


def _try_codec_recovery(text: str) -> str | None:
    """Repair mojibake by reversing a known wrong-decode, if one fits.

    Returns the recovered text when a candidate pair yields a dramatically
    higher Thai ratio than the input; ``None`` when nothing fits (caller
    keeps the original or falls back to OCR).
    """
    original_ratio = _thai_ratio(text)
    for source_codec, target_codec in _CODEC_RECOVERY_PAIRS:
        try:
            recovered = text.encode(source_codec, errors="strict").decode(target_codec, errors="strict")
        except (UnicodeEncodeError, UnicodeDecodeError):
            continue
        recovered_ratio = _thai_ratio(recovered)
        # Only accept when the fix is decisive: the input carried no usable
        # Thai and the round trip restores a substantial share of it.
        if original_ratio < _MAX_THAI_RATIO_FOR_GARBLED and recovered_ratio > 0.15:
            return recovered
    return None


def _repair_or_flag_pdf_text(text: str) -> str:
    """Gate a PDF text layer: repair known mojibake, else demand OCR.

    1. Deterministic codec recovery (e.g. TIS-620 read as mac-roman) when it
       clearly restores Thai text.
    2. Otherwise, when the layer is long, carries (almost) no Thai, and is
       dominated by extended glyphs — the signature of Thai read through a
       broken font mapping — it is garbled and OCR must produce the text.
       Plain-ASCII layers (English documents) pass through untouched.
    """
    if len(text) < _MIN_TEXT_FOR_LANGUAGE_CHECK:
        return text if _has_meaningful_text(text) else ""
    recovered = _try_codec_recovery(text)
    if recovered is not None:
        return recovered
    if _thai_ratio(text) < _MAX_THAI_RATIO_FOR_GARBLED and _non_ascii_ratio(text) >= _MIN_NON_ASCII_RATIO_FOR_GARBLED:
        return ""
    return text


def _extract_html_flavored_xls(path: Path) -> str:
    """Recover an HTML-flavoured ``.xls`` (SpreadsheetML/MIME-HTML) as text.

    Thai government systems routinely export ".xls" files that are actually
    HTML tables with an Excel namespace.  Neither anydoc (``MalformedError``)
    nor MarkItDown parses that combination, so extract the tables directly
    with the same BeautifulSoup toolchain the .html path already uses (F3).
    """
    raw = path.read_text(encoding="utf-8", errors="replace")
    head = raw.lstrip()[:256].casefold()
    if not head.startswith(("<html", "<table", "<!doctype")):
        # A genuine binary .xls must keep its original failure path.
        raise RuntimeError("FILE_TYPE_NOT_SUPPORTED")
    soup = BeautifulSoup(raw, "html.parser")
    tables = soup.find_all("table")
    if not tables:
        return soup.get_text(" ", strip=True)
    lines: list[str] = []
    for table in tables:
        for row in table.find_all("tr"):
            cells = [" ".join(cell.get_text(" ", strip=True).split()) for cell in row.find_all(["td", "th"])]
            if any(cell for cell in cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text(document: Document) -> str:
    path = Path(document.storage_path)
    ext = path.suffix.lower()
    # Plain-text family never needs conversion: anydoc has no format for
    # .txt/.md/.csv/.json, so read them directly.
    if ext in {".txt", ".md", ".csv", ".json"}:
        return _legacy_extract_text(path)
    # anydoc fast path: Rust conversion + per-page OCR chain for scanned or
    # garbled pages, then the Thai-repair gate. HTML family stays legacy
    # (anydoc has no HTML parser); when the wheel is missing the legacy
    # extractors still serve every other format as a safety net.
    primary_error: Exception | None = None  # first converter failure, used by the terminal raise (review info)
    if ext not in {".html", ".htm"}:
        try:
            from .doc_extraction import extract_document_text
            return extract_document_text(document)
        except RuntimeError as exc:
            if str(exc) != "ANYDOC_UNAVAILABLE":
                raise  # OCR_CHAIN_FAILED / TEXT_EXTRACTION_EMPTY are terminal here.
        except Exception as exc:
            # anydoc signals an unreadable workbook (e.g. an HTML-flavoured
            # ".xls") with MalformedError, not RuntimeError.  Keep the rest of
            # the fallback chain alive instead of failing the document.
            primary_error = exc
        # fall through to the legacy extractors when the wheel is missing
    try:
        text = _markitdown_extract(path)
        if ext == ".pdf":
            text = _repair_or_flag_pdf_text(text)
        if ext == ".pdf" and not text:
            raise RuntimeError("OCR_REQUIRED")
        if text:
            return text
        raise RuntimeError("TEXT_EXTRACTION_EMPTY")
    except RuntimeError as exc:
        if str(exc) == "OCR_REQUIRED":
            raise
        primary_error = exc
    except Exception as exc:
        primary_error = exc
    if ext in LEGACY_EXTRACTOR_EXTENSIONS:
        try:
            text = _legacy_extract_text(path)
            if ext == ".pdf":
                text = _repair_or_flag_pdf_text(text)
            if ext == ".pdf" and not text:
                raise RuntimeError("OCR_REQUIRED")
            if text.strip():
                return text
        except RuntimeError:
            raise
        except Exception as exc:
            raise RuntimeError("TEXT_EXTRACTION_FAILED") from exc
    # Last resort for a ".xls" that is really HTML (Thai government exports,
    # F3): both real converters rejected it, but the bytes are a parseable
    # HTML table.
    if ext == ".xls":
        text = _extract_html_flavored_xls(path)
        if text.strip():
            return text
    raise RuntimeError("TEXT_EXTRACTION_FAILED") from primary_error


def split_text(text: str, chunk_size: int, overlap: int) -> list[tuple[int, int, str]]:
    """Split normalized text on word boundaries with bounded overlap."""
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    chunks: list[tuple[int, int, str]] = []
    start = 0
    while start < len(normalized):
        end = min(len(normalized), start + chunk_size)
        if end < len(normalized):
            boundary = normalized.rfind(" ", start + max(1, chunk_size // 2), end)
            if boundary > start:
                end = boundary
        content = normalized[start:end].strip()
        if content:
            chunks.append((start, end, content))
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return chunks


_LEGAL_SECTION_PATTERN = re.compile(
    r"^[ \t]*(มาตรา|ข้อ|หมวด|ส่วนที่|บทเฉพาะกาล|บทนิยาม)\s*([0-9๐-๙]+(?:/[0-9๐-๙]+)?)?\s*(ทวิ|ตรี|จัตวา|เบญจ)?",
    re.MULTILINE,
)
_LEGAL_SECTION_DIGITS = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")


def _legal_section_identity(raw_kind: str, number: str | None, suffix: str | None) -> tuple[str, str | None, str]:
    if not number:
        return raw_kind, None, raw_kind
    normalized_number = number.translate(_LEGAL_SECTION_DIGITS) + (f" {suffix}" if suffix else "")
    return raw_kind, normalized_number, f"{raw_kind} {normalized_number}"


def split_legal_text(text: str, chunk_size: int, overlap: int) -> list[tuple[int, int, str, str | None, str | None, str | None]]:
    """Split on มาตรา/ข้อ/หมวด headings so each chunk keeps its section identity.

    A heading is only recognized at the start of a line, so a mid-sentence
    cross-reference like "ให้เป็นไปตามมาตรา 15" never starts a new section. A
    section body that still exceeds chunk_size falls back to split_text's
    word-boundary sub-split, and every resulting piece inherits that section's
    identity. Text before the first heading is tagged "preamble".
    """
    text = text or ""
    matches = list(_LEGAL_SECTION_PATTERN.finditer(text))
    spans: list[tuple[int, int, str, str | None, str | None]] = []
    if not matches:
        spans.append((0, len(text), "preamble", None, None))
    else:
        if matches[0].start() > 0:
            spans.append((0, matches[0].start(), "preamble", None, None))
        for index, match in enumerate(matches):
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            kind, number, label = _legal_section_identity(match.group(1), match.group(2), match.group(3))
            spans.append((match.start(), end, kind, number, label))
    pieces: list[tuple[int, int, str, str | None, str | None, str | None]] = []
    for span_start, span_end, section_kind, section_number, section_label in spans:
        for offset, sub_end, content in split_text(text[span_start:span_end], chunk_size, overlap):
            pieces.append((span_start + offset, span_start + sub_end, content, section_kind, section_number, section_label))
    return pieces


def replace_document_chunks(db: Session, document: Document, text: str) -> None:
    settings = get_settings()
    db.query(DocumentChunk).filter_by(document_id=document.id).delete(synchronize_session=False)
    splitter = split_legal_text if document.document_type in LEGAL_DOCUMENT_TYPES else split_text
    for index, piece in enumerate(splitter(text, settings.default_chunk_size, settings.default_chunk_overlap)):
        char_start, char_end, content = piece[0], piece[1], piece[2]
        section_kind, section_number, section_label = piece[3:] if len(piece) > 3 else (None, None, None)
        db.add(DocumentChunk(
            document_id=document.id,
            knowledge_base_id=document.knowledge_base_id,
            chunk_index=index,
            content=content,
            content_sha256=hashlib.sha256(content.encode()).hexdigest(),
            char_start=char_start,
            char_end=char_end,
            token_count=len(re.findall(r"\S+", content)),
            section_kind=section_kind,
            section_number=section_number,
            section_label=section_label,
        ))
    db.flush()


def link_provisions_to_chunks(db: Session, document: Document) -> int:
    """Attach matching chunk IDs to each Provision entity so the query-time
    resolver can fetch a provision's exact chunk instead of relying on similarity."""
    if document.document_type not in LEGAL_DOCUMENT_TYPES:
        return 0
    provisions = db.query(Entity).filter(
        Entity.knowledge_base_id == document.knowledge_base_id, Entity.entity_type == "Provision",
        Entity.identity_key.like(f"legal:provision:{document.id}:%"), Entity.deleted_at.is_(None),
    ).all()
    if not provisions:
        return 0
    chunks = db.query(DocumentChunk).filter(
        DocumentChunk.document_id == document.id, DocumentChunk.section_number.is_not(None),
    ).all()
    by_number: dict[str, list[str]] = {}
    for chunk in chunks:
        by_number.setdefault(re.sub(r"\s+", "", chunk.section_number).casefold(), []).append(chunk.id)
    linked = 0
    for provision in provisions:
        number = (provision.attributes or {}).get("provision_number")
        chunk_ids = by_number.get(re.sub(r"\s+", "", str(number)).casefold()) if number else None
        if not chunk_ids:
            continue
        provision.attributes = {**(provision.attributes or {}), "chunk_ids": chunk_ids}
        linked += 1
    db.flush()
    return linked


def embed_document_chunks(db: Session, document_id: str) -> None:
    chunks = db.query(DocumentChunk).filter_by(document_id=document_id).order_by(DocumentChunk.chunk_index).all()
    if not chunks:
        return
    client = OpenRouterClient()
    if not client.embeddings_enabled:
        return
    vectors = client.embed_texts([chunk.content for chunk in chunks])
    for chunk, vector in zip(chunks, vectors, strict=True):
        chunk.embedding = vector
    db.flush()


def queue_embedding_reindex(db: Session, knowledge_base_id: str, force: bool = False) -> int:
    documents = db.query(Document).filter_by(knowledge_base_id=knowledge_base_id, status="completed").filter(Document.deleted_at.is_(None)).all()
    queued = 0
    for document in documents:
        # SQLite represents a JSON null as a non-SQL-NULL value in tests, so
        # inspect the mapped value rather than relying on `IS NOT NULL` alone.
        has_embedding = any(chunk.embedding for chunk in db.query(DocumentChunk).filter_by(document_id=document.id))
        active = db.query(ProcessingJob.id).filter(
            ProcessingJob.document_id == document.id,
            ProcessingJob.job_type == "REINDEX_EMBEDDINGS",
            ProcessingJob.status.in_(["queued", "running"]),
        ).first()
        if (force or not has_embedding) and not active:
            db.add(ProcessingJob(document_id=document.id, knowledge_base_id=knowledge_base_id, job_type="REINDEX_EMBEDDINGS"))
            queued += 1
    db.commit()
    return queued


# Upstream provider failures that mean "retry later with fewer parallel
# calls", observed inside LightRAG track error messages.  Matched
# case-insensitively because the remote relays provider text verbatim.
_BUDGET_EXHAUSTED_MARKERS = ("in_flight_budget_exhausted", "in-flight budget", "payment required", "insufficient credits", "402")


def classify_track_failure(error_detail: str | None) -> str:
    """Map a failed LightRAG track onto a platform error code.

    ``RETRIEVAL_ENGINE_BUDGET_EXHAUSTED`` is transient: the shared provider
    account frees in-flight quota on the scale of minutes, so the worker's
    retry backoff can recover it without operator action (F1).  Everything
    else keeps the existing terminal ``RETRIEVAL_ENGINE_REJECTED``.
    """
    detail = (error_detail or "").casefold()
    if any(marker in detail for marker in _BUDGET_EXHAUSTED_MARKERS):
        return "RETRIEVAL_ENGINE_BUDGET_EXHAUSTED"
    return "RETRIEVAL_ENGINE_REJECTED"


# Follow-up stages run after a document is already searchable and are cheap;
# running them before new full pipelines keeps a legal document's registry
# extraction from queueing minutes behind freshly uploaded bulk work (F4).
_FOLLOW_UP_JOB_TYPES = ("EXTRACT_LEGAL_METADATA", "REINDEX_EMBEDDINGS")


def process_next_job(db: Session) -> bool:
    # Cleanup work (remote-index purges) yields to user-facing jobs: a purge
    # sitting in the queue must not consume a worker slot (or a one-shot
    # /internal/process-next tick) while real documents wait.
    job = db.query(ProcessingJob).filter(
        ProcessingJob.status == "queued", ProcessingJob.next_attempt_at <= datetime.utcnow(),
        ProcessingJob.job_type.in_(_FOLLOW_UP_JOB_TYPES),
    ).order_by(ProcessingJob.created_at).first() or db.query(ProcessingJob).filter(
        ProcessingJob.status == "queued", ProcessingJob.next_attempt_at <= datetime.utcnow(),
        ProcessingJob.job_type != "PURGE_REMOTE_INDEX",
    ).order_by(ProcessingJob.created_at).first() or db.query(ProcessingJob).filter(
        ProcessingJob.status == "queued", ProcessingJob.next_attempt_at <= datetime.utcnow()
    ).order_by(ProcessingJob.created_at).first()
    if not job:
        return False
    if job.job_type == "PURGE_REMOTE_INDEX":
        # 4b: cascade a platform soft-delete into the remote retrieval index
        # so ghost sources stop occupying LightRAG's content-hash namespace.
        # Runs BEFORE the generic cancelled-because-deleted branch below.
        job.status, job.current_stage, job.progress_percent, job.attempt_count = "running", "purging_remote_index", 10, job.attempt_count + 1
        db.commit()
        try:
            doc = db.get(Document, job.document_id)
            if not doc:
                raise RuntimeError("DOCUMENT_NOT_FOUND")
            engine = LightRAGRetrievalEngine()
            if engine.enabled:
                remote = engine.find_document(doc.id, doc.knowledge_base_id)
                if remote:
                    engine.delete_remote_document(remote["id"])
            job.status, job.current_stage, job.progress_percent = "completed", "completed", 100
        except Exception as exc:
            logger.exception("remote index purge failed", extra={"job_id": job.id, "document_id": job.document_id})
            code = str(exc) if str(exc).startswith("RETRIEVAL_ENGINE_") else "REMOTE_INDEX_PURGE_FAILED"
            # Transient engine states (BUSY/TIMEOUT) retry like document
            # processing — a busy remote must not strand a ghost (review info).
            if code in {"RETRIEVAL_ENGINE_BUSY", "RETRIEVAL_ENGINE_TIMEOUT"} and job.attempt_count < MAX_PROCESSING_ATTEMPTS:
                job.status, job.error_code, job.error_message = "queued", code, "Remote engine busy; purge retry scheduled."
                job.next_attempt_at = datetime.utcnow() + timedelta(seconds=processing_retry_delay(job.attempt_count))
            else:
                job.status, job.current_stage, job.error_code, job.error_message = "failed", "failed", "REMOTE_INDEX_PURGE_FAILED", str(exc)[:2000]
        db.commit()
        return True
    if job.job_type == "REBUILD_LEGAL_GRAPH":
        job.status, job.current_stage, job.progress_percent, job.attempt_count = "running", "rebuilding_legal_graph", 10, job.attempt_count + 1
        db.commit()
        try:
            if not job.knowledge_base_id:
                raise RuntimeError("KNOWLEDGE_BASE_NOT_FOUND")
            rebuild_legal_graph(db, job.knowledge_base_id)
            job.status, job.current_stage, job.progress_percent = "completed", "completed", 100
        except Exception as exc:
            logger.exception("legal graph rebuild failed", extra={"job_id": job.id, "knowledge_base_id": job.knowledge_base_id})
            job.status, job.current_stage, job.error_code, job.error_message = "failed", "failed", "LEGAL_GRAPH_REBUILD_FAILED", str(exc)[:2000]
        db.commit()
        return True
    doc = db.get(Document, job.document_id)
    if not doc or doc.deleted_at:
        job.status, job.current_stage, job.error_code = "cancelled", "cancelled", "DOCUMENT_DELETED"
        db.commit()
        return True
    reindex_only = job.job_type == "REINDEX_EMBEDDINGS"
    legal_only = job.job_type == "EXTRACT_LEGAL_METADATA"
    job.status, job.current_stage, job.progress_percent, job.attempt_count = "running", "extracting", 10, job.attempt_count + 1
    # Legal metadata runs as a follow-up job.  It must never move an already
    # searchable document back to "extracting" or re-read the source file.
    if not reindex_only and not legal_only:
        doc.status = "extracting"
    db.commit()
    try:
        if (reindex_only or legal_only) and doc.extracted_text:
            text = doc.extracted_text
        else:
            # anydoc path (extract_text) runs the OCR chain internally per
            # page, so a scanned PDF never surfaces OCR_REQUIRED from here
            # anymore; the code remains for the legacy MarkItDown fallback.
            text = extract_text(doc)
        if not reindex_only:
            doc.extracted_text = text
        if legal_only:
            job.current_stage, job.progress_percent = "legal_extraction", 60
            db.commit()
            deterministic = parse_legal_corpus_metadata(text, doc.title or doc.original_filename)
            # Official corpora with a structured header already carry the
            # version identity and explicit change clauses needed by the legal
            # registry.  Avoid an unnecessary model call for those documents;
            # unknown/unstructured legal documents still retain the existing
            # LLM extraction path as a general fallback.
            extracted = {}
            if deterministic.get("instrument", {}).get("version_role") == "unknown":
                try:
                    extracted = OpenRouterClient().extract_legal_metadata(doc.title or doc.original_filename, text)
                except Exception:
                    # Header/change clauses are sufficient to preserve a traceable
                    # legal instrument even when the optional LLM is unavailable.
                    extracted = {}
            if isinstance(extracted, dict):
                extracted["schema_version"] = 2
                extracted["instrument"] = {**(extracted.get("instrument") or {}), **deterministic["instrument"]}
                extracted["change_events"] = deterministic.get("change_events", []) or extracted.get("change_events", [])
                extracted["amendments"] = deterministic.get("amendments", []) or extracted.get("amendments", [])
                extracted["provenance"] = {**(extracted.get("provenance") or {}), **deterministic.get("provenance", {})}
                doc.legal_metadata = extracted
            else:
                doc.legal_metadata = deterministic
            sync_legal_document_graph(db, doc)
            upsert_legal_instrument(db, doc)
            link_provisions_to_chunks(db, doc)
            resolve_instrument_statuses(db, doc.knowledge_base_id)
            job.status, job.current_stage, job.progress_percent = "completed", "completed", 100
            db.commit()
            return True
        job.current_stage, job.progress_percent = "chunking", 45
        existing_chunks = db.query(DocumentChunk.id, DocumentChunk.section_kind).filter_by(document_id=doc.id).all() if reindex_only else []
        # A forced reindex re-chunks a legal document that predates section-aware
        # splitting (no chunk has a section identity yet); every other reindex
        # keeps its existing chunks and only refreshes embeddings.
        needs_legal_rechunk = doc.document_type in LEGAL_DOCUMENT_TYPES and existing_chunks and not any(section_kind for _, section_kind in existing_chunks)
        if reindex_only and existing_chunks and not needs_legal_rechunk:
            pass
        else:
            replace_document_chunks(db, doc, text)
        db.commit()
        sync_document_metadata_values(db, doc)
        sync_document_metadata_graph(db, doc)
        db.commit()
        job.current_stage, job.progress_percent = "embedding", 60
        embed_document_chunks(db, doc.id)
        db.commit()
        engine = LightRAGRetrievalEngine()
        if engine.enabled and not reindex_only:
            job.current_stage, job.progress_percent = "indexing", 70
            db.commit()
            # Bounded parallelism into the remote pipeline (F1): without this
            # cap a bulk upload fans out into enough concurrent upstream LLM
            # calls to exhaust the shared provider budget.
            with _indexing_semaphore:
                doc.external_engine_id = engine.ingest(doc.id, doc.knowledge_base_id, text, doc.title or doc.original_filename)
                track_id = doc.external_engine_id
                if track_id:
                    deadline = time.monotonic() + get_settings().lightrag_processing_timeout_seconds
                    while time.monotonic() < deadline:
                        track = engine.track_status(track_id)
                        if track["status"] == "processed":
                            break
                        if track["status"] == "failed":
                            raise RuntimeError(classify_track_failure(track.get("error")))
                        job.progress_percent = 85
                        db.commit()
                        time.sleep(2)
                    else:
                        raise RuntimeError("RETRIEVAL_ENGINE_TIMEOUT")
            sync_lightrag_document_graph(db, doc)
        if not reindex_only and not legal_only:
            doc.status = "completed"; doc.indexed_at = datetime.utcnow()
            # A retry that finally succeeded must not keep the failure it
            # recovered from — the document view would show an error code on a
            # healthy document (F2).
            doc.error_code, doc.error_message = None, None
            if doc.document_type in LEGAL_DOCUMENT_TYPES:
                db.add(ProcessingJob(
                    document_id=doc.id,
                    knowledge_base_id=doc.knowledge_base_id,
                    job_type="EXTRACT_LEGAL_METADATA",
                    current_stage="queued",
                ))
        job.status, job.current_stage, job.progress_percent = "completed", "completed", 100
        # Review M2: a retried job must not keep the failure code it recovered
        # from — clear the job-level error the same way the document view does
        # (doc-level clear above covers F2's user-visible case).
        job.error_code, job.error_message = None, None
    except Exception as exc:
        code = (str(exc).partition(":")[0] if str(exc).startswith("OCR_CHAIN_FAILED") else str(exc)) if (str(exc).startswith("OCR_CHAIN_FAILED") or str(exc) in {"OCR_REQUIRED", "OCR_CHAIN_FAILED", "TEXT_EXTRACTION_EMPTY", "FILE_TYPE_NOT_SUPPORTED", "RETRIEVAL_ENGINE_UNAVAILABLE", "RETRIEVAL_ENGINE_REJECTED", "RETRIEVAL_ENGINE_BUDGET_EXHAUSTED", "RETRIEVAL_ENGINE_BUSY", "RETRIEVAL_ENGINE_TIMEOUT", "OPENROUTER_UNAVAILABLE", "OPENROUTER_EMBEDDING_INVALID_RESPONSE", "OPENROUTER_EMBEDDING_DIMENSION_MISMATCH", "OPENROUTER_LLM_INVALID_RESPONSE", "EXTERNAL_OCR_NOT_CONFIGURED", "EXTERNAL_OCR_UNAVAILABLE", "EXTERNAL_OCR_REJECTED", "EXTERNAL_OCR_TIMEOUT", "EXTERNAL_OCR_EMPTY_RESULT", "EXTERNAL_OCR_INVALID_RESPONSE"}) else "TEXT_EXTRACTION_FAILED"
        logger.exception("document processing failed", extra={"document_id": doc.id, "job_id": job.id, "error_code": code})
        message = "The document could not be processed."
        if code == "RETRIEVAL_ENGINE_BUDGET_EXHAUSTED":
            message = "Upstream indexing budget temporarily exhausted; retry scheduled."
        job.error_code, job.error_message = code, message
        attempt_ceiling = MAX_BUDGET_PROCESSING_ATTEMPTS if code == "RETRIEVAL_ENGINE_BUDGET_EXHAUSTED" else MAX_PROCESSING_ATTEMPTS
        if code in TRANSIENT_PROCESSING_ERRORS and job.attempt_count < attempt_ceiling:
            delay = processing_retry_delay(job.attempt_count)
            if code == "RETRIEVAL_ENGINE_BUDGET_EXHAUSTED":
                delay = max(delay, BUDGET_RETRY_DELAY_FLOOR_SECONDS)
            job.status, job.current_stage = "queued", "retry_wait"
            job.next_attempt_at = datetime.utcnow() + timedelta(seconds=delay)
            doc.status, doc.error_code, doc.error_message = "queued", code, "Temporary dependency failure; retry scheduled."
        else:
            if legal_only:
                job.error_message = "Legal metadata extraction failed; the source document remains available."
            elif code == "OCR_REQUIRED":
                doc.status, doc.error_code, doc.error_message = "ocr_required", code, "This PDF has no text layer. OCR is required before it can be searched."
            else:
                doc.status, doc.error_code, doc.error_message = "failed", code, message
            job.status = "failed"
    db.commit()
    return True


def process_next_graph_projection(db: Session) -> bool:
    """Deliver one transactionally-recorded graph event with bounded retry backoff."""
    store = Neo4jGraphStore()
    if not store.enabled:
        return False
    event = db.query(GraphProjectionEvent).filter(
        GraphProjectionEvent.status == "queued", GraphProjectionEvent.next_attempt_at <= datetime.utcnow()
    ).order_by(GraphProjectionEvent.created_at).first()
    if not event:
        return False
    event.status, event.attempt_count = "running", event.attempt_count + 1
    db.commit()
    try:
        if event.event_type == "entity" and event.entity_id:
            entity = db.get(Entity, event.entity_id)
            if entity:
                store.upsert_entity(entity)
        elif event.event_type == "relationship" and event.relationship_id:
            relationship = db.get(Relationship, event.relationship_id)
            if relationship:
                source, target = db.get(Entity, relationship.source_entity_id), db.get(Entity, relationship.target_entity_id)
                if source and target:
                    store.upsert_entity(source)
                    store.upsert_entity(target)
                    store.upsert_relationship(relationship, source, target)
        event.status, event.completed_at, event.last_error = "completed", datetime.utcnow(), None
    except (httpx.HTTPError, RuntimeError) as exc:
        event.status = "queued"
        event.last_error = str(exc)[:2000]
        event.next_attempt_at = datetime.utcnow() + timedelta(seconds=min(60, 2 ** event.attempt_count))
    db.commit()
    return True


def plan_intent(query: str) -> str:
    value = query.lower()
    if any(word in value for word in ["impact", "affected", "ผลกระทบ", "หยุดทำงาน"]): return "impact_analysis"
    if any(word in value for word in ["depend", "relationship", "เชื่อม", "สัมพันธ์"]): return "relationship_lookup"
    if re.search(r"\b[a-z]+[-_]?[0-9]{2,}\b", value): return "entity_lookup"
    return "hybrid_fallback"


def build_retrieval_plan(db: Session, query: str, kb_ids: list[str], max_sources: int, query_filters=None,
                         trace: list[dict] | None = None) -> PlannerDecision:
    """Build a policy-constrained plan; invoke the LLM only when rules are ambiguous."""
    rows = db.query(KnowledgeBase).filter(KnowledgeBase.id.in_(kb_ids)).all() if kb_ids else []
    policy = intersect_policies([policy_from_config(row.retrieval_config) for row in rows])
    decision = rule_plan(query, policy, max_sources)
    requested_from = getattr(query_filters, "published_from", None) if query_filters else None
    requested_to = getattr(query_filters, "published_to", None) if query_filters else None
    as_of_date = getattr(query_filters, "as_of_date", None) if query_filters else None
    include_historical = bool(getattr(query_filters, "include_historical", False)) if query_filters else False
    metadata_filters = dict(getattr(query_filters, "metadata", {}) or {}) if query_filters else {}
    decision.plan = decision.plan.model_copy(update={
        "published_from": requested_from or decision.plan.published_from,
        "published_to": requested_to or decision.plan.published_to,
        "as_of_date": as_of_date or decision.plan.as_of_date,
        "include_historical": include_historical or decision.plan.include_historical,
        "metadata_filters": metadata_filters,
    })
    if decision.ambiguous and policy.planner_llm_fallback:
        try:
            available = [channel.value for channel, enabled in (
                (RetrievalChannel.VECTOR, policy.enable_vector),
                (RetrievalChannel.FULLTEXT, policy.enable_fulltext),
                (RetrievalChannel.GRAPH, policy.enable_graph),
                (RetrievalChannel.LIGHTRAG, policy.enable_lightrag),
            ) if enabled]
            value = OpenRouterClient().plan_retrieval(query, available, decision.plan.max_sources, policy.maximum_graph_depth)
            decision = apply_llm_plan(decision, value, policy, max_sources, query=query)
        except (RuntimeError, ValueError, TypeError) as exc:
            decision.plan = decision.plan.model_copy(update={"planner_source": "rules_fallback", "fallback_reason": str(exc)})
    started_at = time.monotonic()
    legal_context = resolve_legal_context(db, query, kb_ids, decision.plan, policy)
    if legal_context is not None:
        decision.plan = decision.plan.model_copy(update={"legal_context": legal_context})
        _append_retrieval_trace(trace, channel="legal_resolver", system="PostgreSQL legal registry", status="used",
                                started_at=started_at, result_count=len(legal_context.current_version_ids),
                                detail=f"matched={len(legal_context.matched_instrument_ids)} excluded={len(legal_context.excluded_document_ids)}")
    metrics.observe_planner(decision.plan.intent, decision.plan.planner_source, decision.plan.fallback_reason is not None)
    return decision


def _append_retrieval_trace(trace: list[dict] | None, *, channel: str, system: str, status: str,
                            started_at: float, result_count: int = 0, detail: str | None = None) -> None:
    """Record an operator-safe description of an actual retrieval attempt."""
    if trace is None:
        return
    item = {"channel": channel, "system": system, "status": status, "result_count": result_count,
            "duration_ms": round((time.monotonic() - started_at) * 1000),
            # Monotonic time is converted into a safe relative offset before
            # persistence. It lets the trace UI accurately show parallel work.
            "started_at_ms": round(started_at * 1000)}
    if detail:
        item["detail"] = detail
    trace.append(item)
    metrics.observe_retrieval(channel, status, item["duration_ms"])


def query_documents(db: Session, query: str, kb_ids: list[str], limit: int,
                    trace: list[dict] | None = None, plan: RetrievalPlan | None = None,
                    warnings: list[dict] | None = None) -> RetrievalEvidence:
    plan = plan or rule_plan(query, policy_from_config(None), limit).plan
    if plan.metadata_filters:
        metadata_document_ids = _metadata_filter_document_ids(db, kb_ids, plan.metadata_filters)
        plan = plan.model_copy(update={"metadata_document_ids": metadata_document_ids})
        _append_retrieval_trace(trace, channel="metadata_filter", system="PostgreSQL document metadata", status="used",
                                started_at=time.monotonic(), result_count=len(metadata_document_ids),
                                detail=f"exact filter keys: {', '.join(sorted(plan.metadata_filters))}")
    if plan.legal_context and plan.legal_context.ambiguous_context and not plan.include_historical:
        detail = plan.legal_context.ambiguity_reason or "The legal instrument or provision context is ambiguous."
        if warnings is not None:
            warnings.append({
                "code": "AMBIGUOUS_LEGAL_CONTEXT",
                "detail": detail,
                "candidate_instrument_ids": plan.legal_context.candidate_instrument_ids,
            })
        _append_retrieval_trace(trace, channel="legal_resolver", system="PostgreSQL legal registry", status="ambiguous",
                                started_at=time.monotonic(), result_count=0, detail=detail)
        return RetrievalEvidence([], [], [], [])
    channels: list[RetrievalEvidence] = []
    db_channels = {
        RetrievalChannel.EXACT: lambda session: query_exact_documents(session, plan, kb_ids, limit, trace),
        RetrievalChannel.VECTOR: lambda session: query_database_vectors(session, query, kb_ids, limit, trace, plan),
        RetrievalChannel.FULLTEXT: lambda session: query_database_chunks(session, query, kb_ids, limit, trace, plan),
        RetrievalChannel.GRAPH: lambda session: query_database_graph(session, query, kb_ids, limit, trace, plan),
    }
    for channel in (RetrievalChannel.EXACT, RetrievalChannel.VECTOR, RetrievalChannel.FULLTEXT, RetrievalChannel.GRAPH):
        if channel not in plan.channels:
            _append_retrieval_trace(trace, channel=channel.value, system="planner", status="skipped",
                                    started_at=time.monotonic(), detail="not selected by retrieval plan")
    engine = LightRAGRetrievalEngine()
    if RetrievalChannel.LIGHTRAG not in plan.channels:
        _append_retrieval_trace(trace, channel="lightrag", system="planner", status="skipped", started_at=time.monotonic(),
                                detail="not selected by retrieval plan")
    elif not engine.enabled or len(kb_ids) != 1:
        _append_retrieval_trace(trace, channel="lightrag", system="LightRAG", status="skipped", started_at=time.monotonic(),
                                detail="not configured" if not engine.enabled else "requires one Knowledge Base")
    else:
        db_channels[RetrievalChannel.LIGHTRAG] = lambda session: _query_lightrag(session, engine, query, kb_ids, limit, trace, plan)
    futures = {}
    channel_results = {}
    with ThreadPoolExecutor(max_workers=max(1, len(db_channels))) as executor:
        for channel, callback in db_channels.items():
            if channel in plan.channels:
                futures[executor.submit(_run_retrieval_channel, callback, True)] = channel
        for future in as_completed(futures):
            try:
                value = future.result()
                if value is not None:
                    channel_results[futures[future]] = value
            except Exception as exc:
                channel = futures[future]
                _append_retrieval_trace(trace, channel=channel.value, system="executor", status="unavailable",
                                        started_at=time.monotonic(), detail=str(exc))
    channels = [channel_results[channel] for channel in (RetrievalChannel.EXACT, RetrievalChannel.VECTOR, RetrievalChannel.FULLTEXT, RetrievalChannel.GRAPH, RetrievalChannel.LIGHTRAG) if channel in channel_results]

    if plan.legal_context and plan.legal_context.excluded_document_ids and not plan.include_historical:
        excluded = set(plan.legal_context.excluded_document_ids)
        removed, rebuilt = 0, []
        for channel_evidence in channels:
            kept = [source for source in channel_evidence.sources if source.get("document_id") not in excluded]
            removed += len(channel_evidence.sources) - len(kept)
            rebuilt.append(RetrievalEvidence(kept, channel_evidence.entities, channel_evidence.relationships, channel_evidence.paths, channel_evidence.answer))
        channels = rebuilt
        _append_retrieval_trace(trace, channel="legal_validity_filter", system="PostgreSQL legal registry", status="used",
                                started_at=time.monotonic(), result_count=removed,
                                detail=f"removed {removed} superseded/repealed source(s)")

    if plan.legal_context and plan.legal_context.preferred_document_ids and not plan.include_historical:
        preferred = set(plan.legal_context.preferred_document_ids)
        removed, rebuilt = 0, []
        for channel_evidence in channels:
            kept = [source for source in channel_evidence.sources if source.get("document_id") in preferred]
            removed += len(channel_evidence.sources) - len(kept)
            rebuilt.append(RetrievalEvidence(kept, channel_evidence.entities, channel_evidence.relationships,
                                              channel_evidence.paths, channel_evidence.answer))
        channels = rebuilt
        _append_retrieval_trace(trace, channel="legal_context_filter", system="PostgreSQL legal registry", status="used",
                                started_at=time.monotonic(), result_count=removed,
                                detail="preferred current consolidated instrument")

    if plan.intent == "legal_provision" and plan.legal_context and plan.legal_context.provision_refs:
        refs = parse_provision_refs(" ".join(plan.legal_context.provision_refs))
        removed, rebuilt = 0, []
        for channel_evidence in channels:
            kept = [
                source for source in channel_evidence.sources
                if any(
                    provision_number_matches(ref["number"], source.get("section_number"))
                    and (not source.get("section_kind") or source.get("section_kind") == ref["kind"])
                    for ref in refs
                )
            ]
            removed += len(channel_evidence.sources) - len(kept)
            rebuilt.append(RetrievalEvidence(kept, channel_evidence.entities, channel_evidence.relationships,
                                              channel_evidence.paths, channel_evidence.answer))
        channels = rebuilt
        _append_retrieval_trace(trace, channel="provision_filter", system="PostgreSQL legal registry", status="used",
                                started_at=time.monotonic(), result_count=removed,
                                detail="kept only requested provision evidence")

    channels = _prefer_consolidated_body_sources(db, channels, plan, trace)

    document_ids = {source["document_id"] for channel_evidence in channels for source in channel_evidence.sources}
    instruments_by_document = _legal_instruments_by_document(db, document_ids)
    legal_meta = {doc_id: {"authority_level": row.authority_level, "status": row.status, "effective_from": row.effective_from}
                 for doc_id, row in instruments_by_document.items()}
    fused = fuse_evidence(*channels, limit=limit, legal_meta=legal_meta, legal_context=plan.legal_context,
                          authority_weight=plan.authority_weight, status_weight=plan.status_weight, recency_weight=plan.recency_weight)
    fused = _expand_requested_legal_sections(db, fused, plan, trace)
    evidence = rerank_evidence(query, fused, limit, trace, plan.rerank_enabled)
    if evidence.sources is not fused.sources:
        # The cross-encoder reranker has no legal awareness and can drop the
        # current-version guarantee fuse_evidence already applied; restore it
        # from the pre-rerank ranking, still capped at `limit`.
        guaranteed_sources = _apply_current_version_guarantee(evidence.sources, plan.legal_context, limit, fused.sources)
        if guaranteed_sources != evidence.sources:
            sources = [{**item, "citation_id": f"S{index}"} for index, item in enumerate(guaranteed_sources, 1)]
            evidence = RetrievalEvidence(sources, evidence.entities, evidence.relationships, evidence.paths, evidence.answer)

    if instruments_by_document and evidence.sources:
        conflict_warnings = validate_legal_evidence(db, evidence, plan, instruments_by_document)
        if conflict_warnings:
            if warnings is not None:
                warnings.extend(conflict_warnings)
            _append_retrieval_trace(trace, channel="conflict_check", system="Legal registry", status="used",
                                    started_at=time.monotonic(), result_count=len(conflict_warnings),
                                    detail="version/provision conflicts detected")
        _decorate_sources_with_legal_metadata(evidence.sources, instruments_by_document)

    if evidence.sources:
        started_at = time.monotonic()
        try:
            evidence.answer = OpenRouterClient().answer_from_sources(query, evidence.sources)
            _append_retrieval_trace(trace, channel="answer_generation", system="OpenRouter LLM", status="used",
                                    started_at=started_at, result_count=len(evidence.sources), detail="cited answer")
        except RuntimeError as exc:
            evidence.answer = None
            _append_retrieval_trace(trace, channel="answer_generation", system="OpenRouter LLM", status="unavailable",
                                    started_at=started_at, detail=str(exc))
        _validate_answer_citations(evidence, warnings)
    return evidence


def _expand_requested_legal_sections(db: Session, evidence: RetrievalEvidence, plan: RetrievalPlan,
                                     trace: list[dict] | None = None) -> RetrievalEvidence:
    """Replace a matching legal chunk with the complete, bounded provision.

    Legal lists and exceptions commonly cross chunk boundaries.  Passing only
    the top-ranked fragment to answer generation makes the model omit items;
    this deterministic expansion preserves one citation per provision while
    giving the generator every chunk that belongs to it.
    """
    if plan.intent != "legal_provision" or not evidence.sources:
        return evidence
    refs = parse_provision_refs(" ".join(plan.legal_context.provision_refs)) if plan.legal_context else []
    if not refs:
        return evidence
    rebuilt, expanded = [], 0
    seen: set[tuple[str, str, str]] = set()
    for source in evidence.sources:
        document_id, kind, number = source.get("document_id"), source.get("section_kind"), source.get("section_number")
        if not document_id or not kind or not number or not any(
            provision_number_matches(ref["number"], number) and ref["kind"] == kind for ref in refs
        ):
            rebuilt.append(source)
            continue
        key = (document_id, kind, number)
        if key in seen:
            continue
        seen.add(key)
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id, DocumentChunk.section_kind == kind,
            DocumentChunk.section_number == number,
        ).order_by(DocumentChunk.chunk_index).all()
        if not chunks:
            rebuilt.append(source)
            continue
        excerpt = "\n\n".join(chunk.content for chunk in chunks)
        rebuilt.append({**source, "chunk_id": chunks[0].id, "chunk_index": chunks[0].chunk_index,
                        "chunk_ids": [chunk.id for chunk in chunks], "excerpt": excerpt[:12000]})
        expanded += max(0, len(chunks) - 1)
    if expanded:
        _append_retrieval_trace(trace, channel="legal_section_expansion", system="PostgreSQL document chunks", status="used",
                                started_at=time.monotonic(), result_count=expanded,
                                detail="expanded requested provision to all section chunks")
    return RetrievalEvidence([{**source, "citation_id": f"S{index}"} for index, source in enumerate(rebuilt, 1)],
                             evidence.entities, evidence.relationships, evidence.paths, evidence.answer)


def _run_retrieval_channel(callback, needs_db: bool):
    session = SessionLocal() if needs_db else None
    try:
        return callback(session)
    finally:
        if session is not None:
            session.close()


def _query_lightrag(db: Session, engine: LightRAGRetrievalEngine, query: str, kb_ids: list[str], limit: int,
                    trace: list[dict] | None = None, plan: RetrievalPlan | None = None) -> RetrievalEvidence:
    started_at = time.monotonic()
    try:
        value = engine.query(query, kb_ids, limit)
        if plan and _relationship_filter_active(plan):
            source_ids = {source.get("document_id") for source in value.sources if source.get("document_id")}
            allowed = db.query(Document.id).filter(
                Document.id.in_(source_ids), Document.status == "completed", Document.deleted_at.is_(None),
            ) if source_ids else db.query(Document.id).filter(False)
            allowed = _apply_published_filter(allowed, plan)
            allowed = _apply_metadata_filter(allowed, plan)
            allowed_ids = {row[0] for row in allowed.all()}
            value = RetrievalEvidence(
                [source for source in value.sources if source.get("document_id") in allowed_ids],
                value.entities, value.relationships, value.paths, value.answer,
            )
        _append_retrieval_trace(trace, channel="lightrag", system="LightRAG", status="used",
                                started_at=started_at, result_count=len(value.sources), detail="mix retrieval")
        return value
    except RuntimeError as exc:
        _append_retrieval_trace(trace, channel="lightrag", system="LightRAG", status="unavailable",
                                started_at=started_at, detail=str(exc))
        return RetrievalEvidence([], [], [], [])


def _apply_published_filter(rows, plan: RetrievalPlan | None):
    if not plan:
        return rows
    if plan.published_from:
        rows = rows.filter(Document.published_at >= plan.published_from)
    if plan.published_to:
        rows = rows.filter(Document.published_at < plan.published_to)
    return rows


def _metadata_filter_document_ids(db: Session, knowledge_base_ids: list[str], filters: dict[str, str]) -> list[str]:
    """Resolve exact filters through the indexed, filterable metadata projection."""
    if not filters:
        return []
    rows = db.query(Document.id).join(
        DocumentMetadataValue, DocumentMetadataValue.document_id == Document.id,
    ).filter(Document.deleted_at.is_(None))
    if knowledge_base_ids:
        rows = rows.filter(Document.knowledge_base_id.in_(knowledge_base_ids))
    rows = rows.filter(or_(*[
        and_(DocumentMetadataValue.field_key == key, DocumentMetadataValue.value_text == str(value))
        for key, value in filters.items()
    ]))
    # A document must match every requested field. The indexed table has one
    # row per document/field, so grouping avoids returning partial matches.
    rows = rows.group_by(Document.id).having(func.count(func.distinct(DocumentMetadataValue.field_key)) == len(filters))
    return [row[0] for row in rows.all()]


def _apply_metadata_filter(rows, plan: RetrievalPlan | None):
    if plan is None or plan.metadata_document_ids is None:
        return rows
    return rows.filter(Document.id.in_(plan.metadata_document_ids))


def _apply_legal_filter(rows, plan: RetrievalPlan | None):
    """Drop rows for documents the legal resolver marked repealed/superseded at
    the resolved as-of date; a no-op unless the query actually touched a legal KB."""
    if not plan or plan.include_historical or not plan.legal_context or not plan.legal_context.excluded_document_ids:
        return rows
    return rows.filter(Document.id.notin_(plan.legal_context.excluded_document_ids))


def _apply_provision_filter(rows, plan: RetrievalPlan | None):
    """Restrict legal provision retrieval to the requested article/section.

    A bare number such as 6 appears in every amendment file.  Once the legal
    resolver selects an instrument, returning unrelated sections from that
    instrument is still unsafe because the answer model may cite them.
    """
    if not plan or plan.intent != "legal_provision" or not plan.legal_context or not plan.legal_context.provision_refs:
        return rows
    refs = parse_provision_refs(" ".join(plan.legal_context.provision_refs))
    if not refs:
        return rows
    predicates = [
        (DocumentChunk.section_kind == ref["kind"]) & (DocumentChunk.section_number == ref["number"])
        for ref in refs
    ]
    return rows.filter(or_(*predicates))


def _prefer_consolidated_body_sources(db: Session, channels: list[RetrievalEvidence], plan: RetrievalPlan,
                                      trace: list[dict] | None = None) -> list[RetrievalEvidence]:
    """Select the code-body occurrence when a consolidated file also contains
    the enabling Act and amendment footnotes with the same article number.

    The official corpus intentionally preserves those appendices for
    provenance.  They must not compete with the first occurrence in the
    consolidated code body for a normal ``ประมวลกฎหมาย ... มาตรา`` query.
    """
    if not plan.legal_context or not plan.legal_context.preferred_document_ids or not plan.legal_context.provision_refs:
        return channels
    refs = parse_provision_refs(" ".join(plan.legal_context.provision_refs))
    preferred = set(plan.legal_context.preferred_document_ids)
    body_starts: dict[str, int] = {}
    for document_id in preferred:
        value = db.query(func.min(DocumentChunk.chunk_index)).filter(
            DocumentChunk.document_id == document_id, DocumentChunk.section_kind == "หมวด"
        ).scalar()
        if value is not None:
            body_starts[document_id] = int(value)
    if not body_starts:
        return channels

    # Pick one canonical body chunk per requested provision.  We deliberately
    # compare chunk indexes globally across channels because vector, FTS and
    # LightRAG may return the same provision from different routes.
    chosen_indices: dict[tuple[str, str], int] = {}
    for channel in channels:
        for source in channel.sources:
            document_id = source.get("document_id")
            chunk_index = source.get("chunk_index")
            if document_id not in body_starts or chunk_index is None or int(chunk_index) < body_starts[document_id]:
                continue
            for ref in refs:
                if provision_number_matches(ref["number"], source.get("section_number")) and (
                    not source.get("section_kind") or source.get("section_kind") == ref["kind"]
                ):
                    key = (document_id, ref["number"])
                    idx = int(chunk_index)
                    chosen_indices[key] = min(chosen_indices.get(key, idx), idx)

    if not chosen_indices:
        return channels
    rebuilt, removed = [], 0
    # Keep the earliest matching body chunk per requested number.  This also
    # removes the same-number enabling-Act and amendment appendix chunks.
    for channel in channels:
        kept = []
        for source in channel.sources:
            if any(
                source.get("document_id") == document_id
                and provision_number_matches(number, source.get("section_number"))
                and source.get("chunk_index") == chosen_indices.get((document_id, number))
                for document_id, number in chosen_indices
            ):
                kept.append(source)
            elif source.get("document_id") in preferred:
                removed += 1
        rebuilt.append(RetrievalEvidence(kept, channel.entities, channel.relationships, channel.paths, channel.answer))
    _append_retrieval_trace(trace, channel="legal_body_selector", system="PostgreSQL legal registry", status="used",
                            started_at=time.monotonic(), result_count=removed,
                            detail="selected earliest occurrence in consolidated code body")
    return rebuilt


def query_database_vectors(db: Session, query: str, kb_ids: list[str], limit: int,
                           trace: list[dict] | None = None, plan: RetrievalPlan | None = None) -> RetrievalEvidence:
    started_at = time.monotonic()
    if db.get_bind().dialect.name != "postgresql":
        _append_retrieval_trace(trace, channel="semantic_vector", system="PostgreSQL + pgvector", status="skipped",
                                started_at=started_at, detail="PostgreSQL is not the active database")
        return RetrievalEvidence([], [], [], [])
    client = OpenRouterClient()
    if not client.embeddings_enabled:
        _append_retrieval_trace(trace, channel="semantic_vector", system="PostgreSQL + pgvector", status="skipped",
                                started_at=started_at, detail="embedding provider is not configured")
        return RetrievalEvidence([], [], [], [])
    try:
        query_vector = client.embed_texts([query])[0]
    except RuntimeError as exc:
        _append_retrieval_trace(trace, channel="semantic_vector", system="OpenRouter embeddings → PostgreSQL + pgvector",
                                status="unavailable", started_at=started_at, detail=str(exc))
        return RetrievalEvidence([], [], [], [])
    distance = DocumentChunk.embedding.cosine_distance(query_vector)
    rows = db.query(DocumentChunk, Document, distance.label("distance")).join(
        Document, Document.id == DocumentChunk.document_id
    ).filter(Document.status == "completed", Document.deleted_at.is_(None), DocumentChunk.embedding.is_not(None))
    if kb_ids:
        rows = rows.filter(DocumentChunk.knowledge_base_id.in_(kb_ids))
    rows = _apply_published_filter(rows, plan)
    rows = _apply_metadata_filter(rows, plan)
    rows = _apply_legal_filter(rows, plan)
    rows = _apply_provision_filter(rows, plan)
    records = rows.order_by(distance).limit(limit).all()
    sources = [{"citation_id": f"S{i}", "document_id": document.id, "title": document.title,
                "chunk_id": chunk.id, "chunk_index": chunk.chunk_index, "excerpt": chunk.content[:500], "relevance": max(0.0, 1.0 - float(item_distance)),
                "section_kind": chunk.section_kind, "section_number": chunk.section_number, "section_label": chunk.section_label}
               for i, (chunk, document, item_distance) in enumerate(records, 1)]
    _append_retrieval_trace(trace, channel="semantic_vector", system="OpenRouter embeddings → PostgreSQL + pgvector",
                            status="used", started_at=started_at, result_count=len(sources), detail="cosine similarity")
    return RetrievalEvidence(sources, [], [], [])


def query_database_chunks(db: Session, query: str, kb_ids: list[str], limit: int,
                          trace: list[dict] | None = None, plan: RetrievalPlan | None = None) -> RetrievalEvidence:
    started_at = time.monotonic()
    words = [word for word in re.findall(r"[\w-]+", query.lower()) if len(word) > 1]
    base_rows = db.query(DocumentChunk, Document).join(Document, Document.id == DocumentChunk.document_id).filter(
        Document.status == "completed", Document.deleted_at.is_(None)
    )
    if kb_ids:
        base_rows = base_rows.filter(DocumentChunk.knowledge_base_id.in_(kb_ids))
    base_rows = _apply_published_filter(base_rows, plan)
    base_rows = _apply_metadata_filter(base_rows, plan)
    base_rows = _apply_legal_filter(base_rows, plan)
    base_rows = _apply_provision_filter(base_rows, plan)
    rows = base_rows
    metadata_text = cast(Document.metadata_search_text, Text)
    legacy_metadata_text = cast(Document.document_metadata, Text)
    def searchable_metadata_predicate(pattern: str):
        return or_(metadata_text.ilike(pattern), and_(Document.metadata_search_text.is_(None), legacy_metadata_text.ilike(pattern)))
    if words:
        if db.get_bind().dialect.name == "postgresql":
            vector = func.to_tsvector("simple", DocumentChunk.content)
            tsquery = func.websearch_to_tsquery("simple", query)
            rows = rows.filter(or_(vector.op("@@")(tsquery), searchable_metadata_predicate(f"%{query}%"))).order_by(func.ts_rank_cd(vector, tsquery).desc())
        else:
            rows = rows.filter(or_(*[DocumentChunk.content.ilike(f"%{word}%") for word in words[:8]], searchable_metadata_predicate(f"%{query}%")))
    records = rows.limit(limit).all()
    if not records and words and db.get_bind().dialect.name == "postgresql":
        # PostgreSQL simple FTS does not segment Thai. Retain FTS as the first
        # choice, then use an explicit phrase/token fallback in the same scope.
        records = base_rows.filter(or_(*[DocumentChunk.content.ilike(f"%{word}%") for word in words[:8]], *[searchable_metadata_predicate(f"%{word}%") for word in words[:8]])).limit(limit).all()
    sources = [{"citation_id": f"S{i}", "document_id": document.id, "title": document.title,
                "chunk_id": chunk.id, "chunk_index": chunk.chunk_index, "excerpt": chunk.content[:500], "relevance": 1.0 / i,
                "section_kind": chunk.section_kind, "section_number": chunk.section_number, "section_label": chunk.section_label}
               for i, (chunk, document) in enumerate(records, 1)]
    system = "PostgreSQL full-text search" if db.get_bind().dialect.name == "postgresql" else "SQL text fallback"
    _append_retrieval_trace(trace, channel="full_text", system=system, status="used",
                            started_at=started_at, result_count=len(sources), detail="keyword retrieval")
    return RetrievalEvidence(sources, [], [], [])


def query_exact_documents(db: Session, plan: RetrievalPlan, kb_ids: list[str], limit: int,
                          trace: list[dict] | None = None) -> RetrievalEvidence:
    started_at = time.monotonic()
    identifiers = plan.document_identifiers
    if not identifiers:
        _append_retrieval_trace(trace, channel="exact_document", system="PostgreSQL document metadata", status="used",
                                started_at=started_at, detail="no document identifier supplied")
        return RetrievalEvidence([], [], [], [])
    rows = db.query(DocumentChunk, Document).join(Document, Document.id == DocumentChunk.document_id).filter(
        Document.status == "completed", Document.deleted_at.is_(None)
    )
    if kb_ids:
        rows = rows.filter(DocumentChunk.knowledge_base_id.in_(kb_ids))
    rows = _apply_published_filter(rows, plan)
    rows = _apply_metadata_filter(rows, plan)
    rows = _apply_legal_filter(rows, plan)
    rows = _apply_provision_filter(rows, plan)
    predicates = []
    metadata_text = cast(Document.metadata_search_text, Text)
    legacy_metadata_text = cast(Document.document_metadata, Text)
    for identifier in identifiers:
        pattern = f"%{identifier}%"
        predicates.extend((Document.title.ilike(pattern), Document.original_filename.ilike(pattern), DocumentChunk.content.ilike(pattern),
                           metadata_text.ilike(pattern), and_(Document.metadata_search_text.is_(None), legacy_metadata_text.ilike(pattern))))
    records = rows.filter(or_(*predicates)).order_by(DocumentChunk.chunk_index).limit(limit).all()
    sources = [{"citation_id": f"S{i}", "document_id": document.id, "title": document.title,
                "chunk_id": chunk.id, "chunk_index": chunk.chunk_index, "excerpt": chunk.content[:500], "relevance": 1.0 / i,
                "section_kind": chunk.section_kind, "section_number": chunk.section_number, "section_label": chunk.section_label}
               for i, (chunk, document) in enumerate(records, 1)]
    _append_retrieval_trace(trace, channel="exact_document", system="PostgreSQL document metadata", status="used",
                            started_at=started_at, result_count=len(sources), detail="title, filename, and document ID lookup")
    return RetrievalEvidence(sources, [], [], [])


def _verified_relationships(db: Session, kb_ids: list[str]):
    rows = db.query(Relationship).filter(
        Relationship.deleted_at.is_(None),
        (Relationship.is_legal.is_(False)) | (Relationship.review_status == "verified"),
    )
    if kb_ids:
        rows = rows.filter(Relationship.knowledge_base_id.in_(kb_ids))
    return rows


def _global_graph_evidence(db: Session, kb_ids: list[str], limit: int, plan: RetrievalPlan) -> RetrievalEvidence:
    """Return representative cited edges from several connected components."""
    relationship_query = _verified_relationships(db, kb_ids)
    if _relationship_filter_active(plan):
        relationship_query = relationship_query.join(
            RelationshipSource, RelationshipSource.relationship_id == Relationship.id,
        ).join(Document, Document.id == RelationshipSource.document_id).filter(
            Document.status == "completed", Document.deleted_at.is_(None),
        )
        relationship_query = _apply_published_filter(relationship_query, plan)
        relationship_query = _apply_metadata_filter(relationship_query, plan).distinct()
    relationships = relationship_query.order_by(Relationship.source_count.desc(), Relationship.created_at.desc()).limit(max(limit * 10, 50)).all()
    if not relationships:
        return RetrievalEvidence([], [], [], [])
    parent: dict[str, str] = {}
    def find(value: str) -> str:
        parent.setdefault(value, value)
        if parent[value] != value:
            parent[value] = find(parent[value])
        return parent[value]
    def union(first: str, second: str) -> None:
        first, second = find(first), find(second)
        if first != second:
            parent[second] = first
    for relationship in relationships:
        union(relationship.source_entity_id, relationship.target_entity_id)
    selected: list[Relationship] = []
    seen_components: set[str] = set()
    for relationship in relationships:
        component = find(relationship.source_entity_id)
        if component not in seen_components:
            selected.append(relationship); seen_components.add(component)
        if len(selected) >= limit:
            break
    if len(selected) < limit:
        selected.extend(item for item in relationships if item not in selected)
    selected = selected[:limit]
    node_ids = {item.source_entity_id for item in selected} | {item.target_entity_id for item in selected}
    nodes = db.query(Entity).filter(Entity.id.in_(node_ids), Entity.deleted_at.is_(None)).all()
    edges = [{"id": item.id, "source": item.source_entity_id, "target": item.target_entity_id, "type": item.relationship_type} for item in selected]
    sources = relationship_sources(db, selected, plan)[:limit]
    for source in sources:
        source.pop("_relationship_id", None)
    return RetrievalEvidence(sources, [{"id": item.id, "name": item.name, "type": item.entity_type} for item in nodes], edges, edges)


def query_database_graph(db: Session, query: str, kb_ids: list[str], limit: int,
                         trace: list[dict] | None = None, plan: RetrievalPlan | None = None) -> RetrievalEvidence:
    """Use local entity traversal or bounded global graph evidence as planned."""
    started_at = time.monotonic()
    plan = plan or rule_plan(query, policy_from_config(None), limit).plan
    if plan.graph_scope == "global":
        evidence = _global_graph_evidence(db, kb_ids, limit, plan)
        _append_retrieval_trace(trace, channel="graph", system="PostgreSQL graph tables", status="used",
                                started_at=started_at, result_count=len(evidence.sources), detail="global representative relationship evidence")
        return evidence
    subject = next(iter(plan.entity_subjects), query)
    entity = resolve_entity(db, kb_ids, subject)
    if not entity:
        _append_retrieval_trace(trace, channel="graph", system="PostgreSQL graph tables", status="used",
                                started_at=started_at, detail="no scoped entity matched")
        return RetrievalEvidence([], [], [], [])
    store = Neo4jGraphStore()
    accelerator = "postgresql fallback"
    relationship_ids: list[str] = []
    node_ids: list[str] = []
    if store.enabled:
        try:
            result = store.traverse(entity.id, entity.knowledge_base_id, plan.graph_depth, max(limit * 10, 50))
            relationship_ids, node_ids = result.get("relationship_ids", []), result.get("node_ids", [])
            if relationship_ids:
                accelerator = "neo4j accelerator → PostgreSQL evidence"
        except (httpx.HTTPError, RuntimeError, ValueError):
            relationship_ids = []
    if relationship_ids:
        relationships = _verified_relationships(db, [entity.knowledge_base_id]).filter(Relationship.id.in_(relationship_ids)).all()
        node_rows = db.query(Entity).filter(Entity.id.in_(node_ids), Entity.deleted_at.is_(None)).all()
        graph = {"nodes": [{"id": row.id, "name": row.name, "type": row.entity_type} for row in node_rows],
                 "edges": [{"id": row.id, "source": row.source_entity_id, "target": row.target_entity_id, "type": row.relationship_type} for row in relationships]}
    else:
        graph = entity_graph(db, entity, plan.graph_depth)
        relationships = _verified_relationships(db, [entity.knowledge_base_id]).filter(Relationship.id.in_([edge["id"] for edge in graph["edges"]])).all()
    relationships = _filter_relationships_by_plan(db, relationships, plan)
    node_ids = {item.source_entity_id for item in relationships} | {item.target_entity_id for item in relationships}
    node_rows = db.query(Entity).filter(Entity.id.in_(node_ids), Entity.deleted_at.is_(None)).all() if node_ids else []
    graph = {
        "nodes": [{"id": row.id, "name": row.name, "type": row.entity_type} for row in node_rows],
        "edges": [{"id": row.id, "source": row.source_entity_id, "target": row.target_entity_id, "type": row.relationship_type} for row in relationships],
    }
    sources = relationship_sources(db, relationships, plan)[:limit]
    for source in sources:
        source.pop("_relationship_id", None)
    _append_retrieval_trace(trace, channel="graph", system=accelerator, status="used",
                            started_at=started_at, result_count=len(sources), detail="verified/manual bounded traversal")
    return RetrievalEvidence(sources, graph["nodes"], graph["edges"], graph["edges"])


_LEGAL_STATUS_FACTORS = {
    "in_force": 1.0, "amended": 0.85, "not_yet_effective": 0.3, "unknown": 0.5, "superseded": 0.1, "repealed": 0.1,
}


def _legal_instruments_by_document(db: Session, document_ids: set[str]) -> dict[str, LegalInstrument]:
    if not document_ids:
        return {}
    rows = db.query(LegalInstrument).filter(LegalInstrument.document_id.in_(document_ids)).all()
    return {row.document_id: row for row in rows}


def _apply_current_version_guarantee(kept_sources: list[dict], legal_context: LegalContext | None, limit: int,
                                     pool: list[dict]) -> list[dict]:
    """`kept_sources` is already ranked and capped at `limit`. If a document the
    resolver marked as a current legal version is missing from it, backfill the
    first matching source found in `pool` (a superset, e.g. the pre-rerank
    ranking) and trim back down to `limit` -- the result never exceeds `limit`
    even when several documents need backfilling. A no-op without legal_context."""
    if not legal_context or not legal_context.current_version_ids:
        return kept_sources
    present = {source["document_id"] for source in kept_sources}
    missing = [doc_id for doc_id in legal_context.current_version_ids if doc_id not in present]
    if not missing:
        return kept_sources
    missing_set, seen, backfill = set(missing), set(), []
    for source in pool:
        document_id = source["document_id"]
        if document_id in missing_set and document_id not in seen:
            seen.add(document_id)
            backfill.append(source)
    backfill = backfill[:limit]
    return kept_sources[:max(0, limit - len(backfill))] + backfill


def fuse_evidence(*channels: RetrievalEvidence, limit: int, legal_meta: dict[str, dict] | None = None,
                  legal_context: LegalContext | None = None, authority_weight: float = 0.0,
                  status_weight: float = 0.0, recency_weight: float = 0.0) -> RetrievalEvidence:
    """Use reciprocal-rank fusion across semantic, FTS, and graph channels, then
    apply an authority/status/recency boost for any document with a legal
    registry entry. A document absent from legal_meta gets boost 0, so a
    Knowledge Base with no legal instruments scores exactly as before."""
    candidates: dict[str, dict] = {}
    for channel in channels:
        for rank, source in enumerate(channel.sources, 1):
            # A single document can contribute several relevant passages. Keep
            # each chunk as independent evidence so answer generation receives
            # the passage that actually addresses the question.
            key = f"{source['document_id']}:{source['chunk_id']}"
            score = 1.0 / (60 + rank)
            candidate = candidates.setdefault(key, {"score": 0.0, "source": source})
            candidate["score"] += score
            if source.get("relevance", 0.0) > candidate["source"].get("relevance", 0.0):
                candidate["source"] = source

    legal_meta = legal_meta or {}
    dated_documents = sorted(
        (doc_id for doc_id, meta in legal_meta.items() if meta.get("effective_from")),
        key=lambda doc_id: legal_meta[doc_id]["effective_from"], reverse=True,
    )
    recency_rank = {doc_id: 1.0 - (index / max(len(dated_documents) - 1, 1)) for index, doc_id in enumerate(dated_documents)}
    for candidate in candidates.values():
        meta = legal_meta.get(candidate["source"]["document_id"])
        if not meta:
            continue
        boost = (
            authority_weight * (meta["authority_level"] / 100)
            + status_weight * _LEGAL_STATUS_FACTORS.get(meta["status"], 0.5)
            + recency_weight * recency_rank.get(candidate["source"]["document_id"], 0.5)
        )
        candidate["score"] *= (1 + boost)

    ranked = sorted(candidates.values(), key=lambda item: item["score"], reverse=True)
    ranked_sources = [{**item["source"], "relevance": round(item["score"], 6)} for item in ranked]

    # Guaranteed representation: a "current version" document must appear even
    # if raw similarity pushed it out of the top-`limit`, so an amended act
    # never gets silently displaced by an unrelated FAQ hit. The result never
    # exceeds `limit`, even when several documents need backfilling.
    ordered_sources = _apply_current_version_guarantee(ranked_sources[:limit], legal_context, limit, ranked_sources)
    sources = [{**item, "citation_id": f"S{index}"} for index, item in enumerate(ordered_sources, 1)]
    answer = next((channel.answer for channel in channels if channel.answer), None)
    entities = [item for channel in channels for item in channel.entities]
    relationships = [item for channel in channels for item in channel.relationships]
    paths = [item for channel in channels for item in channel.paths]
    return RetrievalEvidence(sources, entities, relationships, paths, answer)


def rerank_evidence(query: str, evidence: RetrievalEvidence, limit: int,
                    trace: list[dict] | None = None, enabled: bool = True) -> RetrievalEvidence:
    """Optionally apply a cross-encoder reranker after deterministic fusion."""
    client = OpenRouterClient()
    if not enabled or not client.reranker_enabled or not evidence.sources:
        _append_retrieval_trace(trace, channel="rerank", system="OpenRouter reranker", status="skipped", started_at=time.monotonic(),
                                detail="disabled by retrieval policy" if not enabled else ("not configured" if not client.reranker_enabled else "no candidates"))
        return evidence
    started_at = time.monotonic()
    candidates = evidence.sources[:get_settings().rerank_candidate_limit]
    try:
        ranked = client.rerank(query, [item["excerpt"] for item in candidates], limit)
    except RuntimeError as exc:
        _append_retrieval_trace(trace, channel="rerank", system="OpenRouter reranker", status="unavailable",
                                started_at=started_at, detail=str(exc))
        return evidence
    ordered = []
    for index, relevance in ranked:
        if 0 <= index < len(candidates):
            ordered.append({**candidates[index], "relevance": relevance})
    if not ordered:
        _append_retrieval_trace(trace, channel="rerank", system="OpenRouter reranker", status="used",
                                started_at=started_at, detail="no ranked candidates returned")
        return evidence
    sources = [{**source, "citation_id": f"S{index}"} for index, source in enumerate(ordered[:limit], 1)]
    _append_retrieval_trace(trace, channel="rerank", system="OpenRouter reranker", status="used",
                            started_at=started_at, result_count=len(sources), detail="cross-encoder rerank")
    return RetrievalEvidence(sources, evidence.entities, evidence.relationships, evidence.paths, evidence.answer)


_LEGAL_STATUS_LABELS_TH = {
    "in_force": "บังคับใช้", "amended": "แก้ไขเพิ่มเติมแล้ว", "superseded": "ถูกแทนที่",
    "repealed": "ถูกยกเลิก", "not_yet_effective": "ยังไม่มีผลบังคับใช้", "unknown": "ไม่ทราบสถานะ",
}


def validate_legal_evidence(db: Session, evidence: RetrievalEvidence, plan: RetrievalPlan | None,
                            instruments_by_document: dict[str, LegalInstrument]) -> list[dict]:
    """Detect same-provision version conflicts and provision-level overrides in
    the fused evidence before the LLM sees it; mutates evidence.sources in place."""
    if not evidence.sources or not plan:
        return []
    warnings: list[dict] = []

    instrument_ids = [row.id for row in instruments_by_document.values()]
    relations = db.query(LegalInstrumentRelation).filter(
        LegalInstrumentRelation.target_instrument_id.in_(instrument_ids), LegalInstrumentRelation.review_status == "verified",
        LegalInstrumentRelation.target_provision.is_not(None), LegalInstrumentRelation.relation.in_(("REPEALS", "SUPERSEDES", "AMENDS")),
    ).all() if instrument_ids else []
    relations_by_target: dict[str, list[LegalInstrumentRelation]] = {}
    for relation in relations:
        relations_by_target.setdefault(relation.target_instrument_id, []).append(relation)

    def overriding_relation(instrument: LegalInstrument, section_number: str) -> LegalInstrumentRelation | None:
        return next((relation for relation in relations_by_target.get(instrument.id, [])
                    if provision_number_matches(relation.target_provision, section_number)), None)

    # 1. Same family + same section kind + same provision number, appearing in
    # MORE THAN ONE DOCUMENT: a verified provision-level edge is the strongest
    # signal of which one is current; fall back to legal_context, then to
    # whichever source ranked first. section_kind is part of the key because a
    # bare number collides across heading kinds (มาตรา 2 and หมวด 2 are not the
    # same provision); requiring >1 document_id avoids treating a long
    # provision's own sub-split chunks (same document) as a version conflict.
    groups: dict[tuple, list[dict]] = {}
    for source in evidence.sources:
        instrument = instruments_by_document.get(source["document_id"])
        section_number = source.get("section_number")
        if not instrument or not instrument.family_id or not section_number:
            continue
        groups.setdefault((instrument.family_id, source.get("section_kind"), section_number), []).append(source)
    current_ids = set(plan.legal_context.current_version_ids) if plan.legal_context else set()
    to_remove: set[int] = set()
    for (_, _, section_number), group in groups.items():
        group_document_ids = {source["document_id"] for source in group}
        if len(group_document_ids) < 2:
            continue
        overridden_ids = {id(source) for source in group
                          if overriding_relation(instruments_by_document[source["document_id"]], section_number)}
        current_in_group = group_document_ids & current_ids
        if overridden_ids and len(overridden_ids) < len(group):
            keep_candidates = [source for source in group if id(source) not in overridden_ids]
        elif current_in_group and current_in_group != group_document_ids:
            # A strict subset is current; a full or empty overlap is ambiguous.
            keep_candidates = [source for source in group if source["document_id"] in current_in_group]
        elif plan.legal_context is not None:
            # The resolver evaluated this Knowledge Base and could not
            # disambiguate without a reviewed relationship; trust that
            # ambiguity rather than guessing which version to drop.
            continue
        else:
            keep_candidates = group
        keep = keep_candidates[0]
        for source in group:
            if source is keep:
                continue
            if not plan.include_historical:
                to_remove.add(id(source))
            warnings.append({
                "code": "SUPERSEDED_VERSION_REMOVED" if not plan.include_historical else "SUPERSEDED_VERSION_PRESENT",
                "detail": f"{source.get('title')} {section_number} is superseded by the current version",
            })
    if to_remove:
        evidence.sources[:] = [source for source in evidence.sources if id(source) not in to_remove]

    # 2. Provision-level repeal/supersede/amend that still targets a chunk left
    # in evidence (e.g. the whole instrument otherwise remains in force).
    for source in evidence.sources:
        instrument = instruments_by_document.get(source["document_id"])
        section_number = source.get("section_number")
        if not instrument or not section_number:
            continue
        relation = overriding_relation(instrument, section_number)
        if relation and relation.relation in {"REPEALS", "SUPERSEDES"}:
            warnings.append({
                "code": "PROVISION_REPEALED",
                "detail": f"{instrument.official_title or source.get('title')} {section_number} was {relation.relation.lower()}; verify against the amending instrument",
            })
        elif relation and relation.relation == "AMENDS":
            warnings.append({
                "code": "PROVISION_AMENDED",
                "detail": f"{instrument.official_title or source.get('title')} {section_number} was amended; verify against the amending instrument",
            })

    # 3. Unverified validity: source document has no resolvable status.
    for source in evidence.sources:
        instrument = instruments_by_document.get(source["document_id"])
        if instrument and instrument.status == "unknown":
            warnings.append({
                "code": "UNVERIFIED_VALIDITY",
                "detail": f"{instrument.official_title or source.get('title')} has no confirmed effective date",
            })
    return warnings


def _decorate_sources_with_legal_metadata(sources: list[dict], instruments_by_document: dict[str, LegalInstrument]) -> None:
    """Attach version/status metadata to each source so citations and the LLM
    prompt can tell an in-force provision from a repealed or amended one."""
    for source in sources:
        instrument = instruments_by_document.get(source["document_id"])
        if not instrument:
            continue
        source["document_status"] = instrument.status
        source["authority_level"] = instrument.authority_level
        source["kind"] = instrument.kind
        source["version_label"] = instrument.version_label
        source["effective_from"] = instrument.effective_from.isoformat() if instrument.effective_from else None
        source["effective_to"] = instrument.effective_to.isoformat() if instrument.effective_to else None
        source["version_date"] = instrument.version_date.isoformat() if instrument.version_date else None
        source["legal_work_key"] = instrument.legal_work_key
        source["document_class"] = instrument.document_class
        source["version_role"] = instrument.version_role
        source["source_uri"] = instrument.source_uri
        source["source_reference"] = instrument.source_reference
        source["provenance"] = {"origin": "legal_registry", "review_status": instrument.review_status,
                                 "document_id": instrument.document_id, "source_uri": instrument.source_uri,
                                 "source_reference": instrument.source_reference}
        # An unconfirmed status is informational (F6): the document is indexed
        # and its content is evidence.  Only a *known* adverse status belongs
        # in the answer prompt; "ไม่ทราบสถานะ" made the generator refuse to
        # answer and the fail-closed citation check then dropped every source.
        status = instrument.status
        if status == "unknown":
            status = None
        parts = [instrument.official_title or source.get("title")]
        if status:
            parts.append(f"สถานะ: {_LEGAL_STATUS_LABELS_TH.get(status, status)}")
        if source.get("section_label"):
            parts.append(source["section_label"])
        if instrument.effective_from:
            parts.append(f"มีผล: {instrument.effective_from.isoformat()}")
        source["legal_label"] = " | ".join(str(part) for part in parts if part)


_CITATION_RE = re.compile(r"\[(S\d+)\]")


def _validate_answer_citations(evidence: RetrievalEvidence, warnings: list[dict] | None = None) -> None:
    """Keep only evidence IDs actually cited by the generated answer.

    Retrieval returns a candidate pool, while the answer is the user-visible
    claim set.  Persisting all candidates as claim citations made unrelated
    chunks look like proof.  A missing/unknown citation is fail-closed on the
    CLAIMS: the generated answer is dropped because none of it is verifiable,
    but the retrieved evidence itself is kept (F6) — the answer engine simply
    falling back to the evidence listing, the same path an unavailable LLM
    already takes, instead of discarding real evidence the user asked for.
    """
    answer = (evidence.answer or "").strip()
    if not answer:
        return
    cited_ids = set(_CITATION_RE.findall(answer))
    available = {str(source.get("citation_id")): source for source in evidence.sources}
    valid_ids = cited_ids & available.keys()
    unknown_ids = sorted(cited_ids - available.keys())
    if not valid_ids:
        if warnings is not None:
            warnings.append({
                "code": "ANSWER_CITATIONS_MISSING",
                "detail": "The generated answer did not cite a valid supplied evidence ID.",
                "unknown_citation_ids": unknown_ids,
            })
        evidence.answer = None
        return
    if unknown_ids:
        if warnings is not None:
            warnings.append({
                "code": "ANSWER_CITATION_ID_INVALID",
                "detail": "The generated answer referenced an evidence ID that was not supplied.",
                "unknown_citation_ids": unknown_ids,
            })
        # Fail-closed on claims (answer dropped), but keep the evidence the
        # generator DID verify: citing [S1] + hallucinated [S99] narrows the
        # pool to S1 rather than re-listing every retrieved candidate (M1).
        evidence.answer = None
        evidence.sources = [available[citation_id] for citation_id in sorted(valid_ids, key=lambda value: int(value[1:]))]
        return
    evidence.sources = [available[citation_id] for citation_id in sorted(valid_ids, key=lambda value: int(value[1:]))]


def _render_citation_details(sources: list[dict]) -> str:
    lines = []
    for source in sources:
        label = source.get("title") or "เอกสาร"
        if source.get("section_label"):
            label = f"{label}, {source['section_label']}"
        uri = source.get("source_uri")
        detail = f"[{source.get('citation_id')}] {label}"
        if uri:
            detail += f" — {uri}"
        lines.append(detail)
    return "\n".join(lines)


def compose_cited_answer(evidence: RetrievalEvidence, warnings: list[dict] | None = None) -> str:
    if warnings and any(item.get("code") == "AMBIGUOUS_LEGAL_CONTEXT" for item in warnings):
        return "ยังไม่สามารถตอบได้อย่างปลอดภัย เนื่องจากพบหลายฉบับที่อาจตรงกับมาตราที่ระบุ โปรดระบุชื่อฉบับกฎหมายหรือวันที่มีผลบังคับใช้ให้ชัดเจน"
    # F6: an answer dropped by the citation gate still has its evidence — the
    # empty-answer path below lists it instead of denying everything.
    if not evidence.sources:
        return "ไม่พบหลักฐานที่เพียงพอในคลังความรู้ที่เลือก จึงไม่ควรสรุปข้อเท็จจริงเพิ่มเติมจากคำขอนี้"
    citations = " ".join(f"[{source['citation_id']}]" for source in evidence.sources)
    answer = (evidence.answer or "").strip()
    details = _render_citation_details(evidence.sources)
    latest_assumption = any(source.get("version_role") == "latest_consolidated" for source in evidence.sources)
    # The consolidated-version prefix asserts THIS ANSWER follows the latest
    # in-force text — only meaningful with a generated answer present (info#1).
    prefix = ("คำตอบนี้อ้างอิงฉบับปรับปรุงล่าสุดที่มีผลบังคับใช้ในคลังข้อมูล\n"
              if latest_assumption and answer else "")
    if not answer:
        return f"พบหลักฐานที่เกี่ยวข้อง {citations}\n\nแหล่งอ้างอิง:\n{details}"
    if not re.search(r"\[S\d+\]", answer):
        answer = f"{answer}\n\nแหล่งอ้างอิง: {citations}"
    return f"{prefix}{answer}\n\nรายละเอียดแหล่งอ้างอิง:\n{details}"


def build_query_result(db: Session, query: str, kb_ids: list[str], max_sources: int, token_id: str | None = None, query_filters=None) -> dict:
    current_shortcut_allowed = allows_default_current_direct_path(query, query_filters)
    registry_shortcut_allowed = allows_direct_registry_path(query_filters)
    if is_document_inventory_query(query):
        return build_document_inventory_result(db, query, kb_ids, token_id=token_id, query_filters=query_filters,
                                               include_documents=True, max_documents=500)
    if current_shortcut_allowed and is_legal_metadata_lookup(query):
        return build_legal_metadata_result(db, query, kb_ids, token_id=token_id)
    if registry_shortcut_allowed and ("คำพิพากษา" in (query or "").casefold() or "ฎีกา" in (query or "").casefold()) and not has_court_decision_evidence(db, kb_ids):
        return _persist_scope_gap_result(db, query, kb_ids, token_id=token_id)
    if current_shortcut_allowed and is_legal_commencement_lookup(query):
        return build_legal_commencement_result(db, query, kb_ids, token_id=token_id)
    if current_shortcut_allowed and is_legal_document_copy_lookup(query):
        return build_legal_document_copy_result(db, query, kb_ids, token_id=token_id)
    if registry_shortcut_allowed and is_legal_effective_rule_lookup(query):
        return build_legal_effective_rule_result(db, query, kb_ids, token_id=token_id)
    if registry_shortcut_allowed and is_legal_provenance_lookup(query):
        return build_legal_provenance_result(db, query, kb_ids, token_id=token_id)
    if current_shortcut_allowed and parse_provision_refs(query) and not _requested_amendment_number(query):
        current = _sole_latest_legal_instrument(db, kb_ids)
        if current:
            return build_default_current_provision_result(db, query, kb_ids, token_id=token_id)
    if current_shortcut_allowed and is_default_current_legal_lookup(query):
        return build_default_current_legal_result(db, query, kb_ids, token_id=token_id)
    retrieval_trace: list[dict] = []
    legal_warnings: list[dict] = []
    decision = build_retrieval_plan(db, query, kb_ids, max_sources, query_filters, retrieval_trace)
    evidence = query_documents(db, query, kb_ids, decision.plan.max_sources, retrieval_trace, decision.plan, legal_warnings)
    intent = decision.plan.intent
    answer = compose_cited_answer(evidence, legal_warnings)
    if not evidence.sources and decision.plan.legal_context is not None:
        answer = legal_scope_gap_response(query)
    legal_context = decision.plan.legal_context.model_dump(mode="json") if decision.plan.legal_context else None
    legal_status = "not_applicable"
    if legal_context is not None:
        legal_status = "insufficient" if not evidence.sources else ("partial" if legal_warnings else "verified")
    answer_contract = {"status": legal_status, "as_of_date": decision.plan.as_of_date.isoformat() if decision.plan.as_of_date else None,
                       "warnings": legal_warnings, "citation_required": bool(evidence.sources),
                       "claim_citations": [source.get("citation_id") for source in evidence.sources]}
    metrics.observe_query_outcome("insufficient_evidence" if not evidence.sources else "evidence_found")
    result = {"status": "success", "result_id": "", "answer": answer,
              "insufficient_evidence": not bool(evidence.sources), "entities": evidence.entities,
              "relationships": evidence.relationships, "paths": evidence.paths, "sources": evidence.sources,
              "warnings": legal_warnings, "metadata": {"knowledge_base_ids": kb_ids, "retrieval_strategy": intent,
                                                "retrieval_plan": decision.plan.model_dump(mode="json"),
                                                "planner_policy_version": decision.policy_version,
                                                "retrieval_trace": retrieval_trace,
                                                "legal_context": legal_context,
                                                "answer_contract": answer_contract,
                                                # Safe observability summaries. Full prompts, provider
                                                # payloads, and document bodies are never persisted.
                                                "query_preview": query[:500] if get_settings().log_query_text else None,
                                                "query_length": len(query),
                                                "query_sha256": hashlib.sha256(query.encode("utf-8")).hexdigest(),
                                                "answer_preview": answer[:800] if get_settings().log_query_text else None,
                                                "citation_ids": [source.get("citation_id") for source in evidence.sources],
                                                "filter_summary": query_filters.model_dump(mode="json") if query_filters else {},
                                                "response_summary": {"status": "success", "insufficient_evidence": not bool(evidence.sources),
                                                                     "source_count": len(evidence.sources), "entity_count": len(evidence.entities),
                                                                     "relationship_count": len(evidence.relationships)}}}
    saved = QueryResult(token_key_id=token_id, result_json=result, expires_at=datetime.utcnow() + timedelta(minutes=30))
    db.add(saved); db.flush(); result["result_id"] = saved.id; saved.result_json = result; db.commit()
    return result
